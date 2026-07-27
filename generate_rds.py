#!/usr/bin/env python3
"""
Generate comprehensive RDS (Requirements and Design Specification) document
for CAMS - Clinic Appointment Management System.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val) if val else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_para(text, bold=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

# ══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TÀI LIỆU ĐẶC TẢ YÊU CẦU VÀ THIẾT KẾ HỆ THỐNG')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Requirements & Design Specification (RDS)')
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    ('Dự án:', 'Clinic Appointment Management System (CAMS)'),
    ('Phiên bản:', '1.0'),
    ('Ngày tạo:', datetime.date.today().strftime('%d/%m/%Y')),
    ('Công nghệ:', 'Java Servlet + JSP, SQL Server'),
    ('Mô tả:', 'Hệ thống quản lý đặt lịch và khám thai sản'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' ')
    r1.font.bold = True
    r1.font.size = Pt(13)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(value)
    r2.font.size = Pt(13)
    r2.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ══════════════════════════════════════════════════════════════════════
add_heading('MỤC LỤC', 1)
toc_items = [
    'PHẦN I: TỔNG QUAN HỆ THỐNG',
    '  1.1. Giới thiệu',
    '  1.2. Phạm vi hệ thống',
    '  1.3. Các tác nhân (Actors)',
    'PHẦN II: KIẾN TRÚC HỆ THỐNG',
    '  2.1. Kiến trúc tổng thể',
    '  2.2. Cấu trúc thư mục',
    '  2.3. Công nghệ sử dụng',
    'PHẦN III: CƠ SỞ DỮ LIỆU',
    '  3.1. Danh sách bảng',
    '  3.2. Chi tiết từng bảng',
    '  3.3. Mối quan hệ giữa các bảng',
    'PHẦN IV: USE CASE CHI TIẾT',
    '  4.1. Use Case Diagram',
    '  4.2. Patient Use Cases',
    '  4.3. Staff Use Cases',
    '  4.4. Doctor Use Cases',
    '  4.5. Sonographer Use Cases',
    '  4.6. Manager Use Cases',
    '  4.7. Admin Use Cases',
    'PHẦN V: LUỒNG XỬ LÝ NGHIỆP VỤ',
    '  5.1. Luồng đặt lịch → hoàn tất khám',
    '  5.2. Luồng siêu âm',
    '  5.3. Luồng thanh toán',
    '  5.4. Luồng huỷ/đổi lịch',
    '  5.5. Luồng hoàn tiền & không đến khám',
    'PHẦN VI: VALIDATE LOGIC VÀ RÀNG BUỘC',
    '  6.1. Ma trận trạng thái lịch hẹn',
    '  6.2. Giai đoạn khám (ExamStage)',
    '  6.3. Điều kiện chặn/mở cho từng action',
    '  6.4. Validate đặt lịch',
    '  6.5. Validate thanh toán',
    '  6.6. Validate siêu âm',
    'PHẦN VII: BẢO MẬT VÀ PHÂN QUYỀN',
    'PHẦN VIII: QUY TẮC NGHIỆP VỤ ĐÃ CHỐT',
    'PHẦN IX: API / SERVLET REFERENCE',
]
for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    if not item.startswith('  '):
        r.font.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN I: TỔNG QUAN HỆ THỐNG
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN I: TỔNG QUAN HỆ THỐNG', 1)

add_heading('1.1. Giới thiệu', 2)
add_para('CAMS (Clinic Appointment Management System) là hệ thống quản lý đặt lịch và khám thai sản, '
         'phục vụ phòng khám sản phụ khoa. Hệ thống hỗ trợ toàn bộ quy trình từ đặt lịch, '
         'tiếp nhận, khám lâm sàng, chỉ định siêu âm, phân tích AI, kê đơn, thanh toán đến chốt hồ sơ bệnh án.')
add_para('Hệ thống chạy trên nền tảng Java Servlet + JSP, sử dụng SQL Server làm cơ sở dữ liệu, '
         'giao diện tiếng Việt, chạy localhost:8080.')

add_heading('1.2. Phạm vi hệ thống', 2)
add_para('Hệ thống bao gồm các chức năng chính:')
scope_items = [
    'Quản lý đặt lịch khám (Online bởi bệnh nhân, Manual bởi lễ tân)',
    'Quản lý hàng đợi tiếp nhận (Reception Queue)',
    'Quản lý khám lâm sàng và hồ sơ bệnh án',
    'Quản lý chỉ định siêu âm và tích hợp AI Engine phân tích ảnh',
    'Quản lý kê đơn thuốc',
    'Quản lý thanh toán tiền mặt (PRE_EXAM, POST_EXAM, PRESCRIPTION)',
    'Quản lý hoàn tiền (Refund)',
    'Quản lý danh mục dịch vụ y tế và lịch sử điều chỉnh giá',
    'Quản lý danh mục thuốc',
    'Quản lý lịch làm việc bác sĩ (ca/shift)',
    'Quản lý người dùng và phân quyền (RBAC)',
    'Báo cáo doanh thu và thống kê dịch vụ',
    'Thông báo (Notification) cho các sự kiện',
    'Audit log toàn bộ thao tác quan trọng',
]
for item in scope_items:
    add_bullet(item)

add_para('')
add_para('Các giới hạn:', bold=True)
limits = [
    'KHÔNG có cổng thanh toán điện tử. Bệnh nhân trả tiền mặt tại quầy.',
    'Bệnh nhân không bao giờ được tự đánh dấu đã thanh toán.',
    'Bác sĩ lâm sàng không sửa được kết luận của bác sĩ siêu âm.',
    'Staff không sửa được chẩn đoán hay đơn thuốc.',
    'Không có trạng thái nào được phép khiến một ca treo vĩnh viễn.',
]
for item in limits:
    add_bullet(item)

add_heading('1.3. Các tác nhân (Actors)', 2)
add_table(
    ['Actor', 'Vai trò', 'Mô tả'],
    [
        ['Admin', 'Quản trị viên', 'Quản lý người dùng, phân quyền, cấu hình hệ thống, quản lý thuốc và dịch vụ, xem audit log'],
        ['Manager', 'Quản lý phòng khám', 'Quản lý dịch vụ y tế & giá, lịch sử giá, doanh thu, thống kê, duyệt lịch làm việc, quản lý bác sĩ, quản lý khung giờ (time slots)'],
        ['Patient', 'Bệnh nhân', 'Đặt lịch, xem lịch sử khám, xem hồ sơ bệnh án (đã chốt), thanh toán, đổi/huỷ lịch, đánh giá'],
        ['Doctor', 'Bác sĩ lâm sàng', 'Tiếp nhận ca, khám lâm sàng, chỉ định siêu âm, chẩn đoán & kê đơn, chốt hồ sơ, quản lý lịch làm việc cá nhân'],
        ['Sonographer', 'Bác sĩ siêu âm', 'Tiếp nhận ca siêu âm, chụp & tải ảnh, phân tích AI, vẽ thủ công, ký & công bố kết quả'],
        ['Staff', 'Nhân viên lễ tân', 'Quản lý hàng đợi, duyệt lịch, check-in, thu tiền mặt, đánh dấu ưu tiên, tạo lịch thủ công, huỷ lịch, đánh dấu không đến'],
    ],
    [3, 4, 10]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN II: KIẾN TRÚC HỆ THỐNG
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN II: KIẾN TRÚC HỆ THỐNG', 1)

add_heading('2.1. Kiến trúc tổng thể', 2)
add_para('Hệ thống tuân theo kiến trúc Model-View-Controller (MVC) truyền thống với Java Servlet:')
arch_items = [
    'View (JSP): Giao diện người dùng, hiển thị dữ liệu, thu thập input từ form',
    'Controller (Servlet): Nhận HTTP request, gọi Service xử lý nghiệp vụ, forward/redirect tới JSP',
    'Service (Service Layer): Chứa toàn bộ logic nghiệp vụ, validation, orchestration',
    'DAO (Data Access Object): Truy vấn database, trả về Model objects',
    'Model (POJO): Đối tượng dữ liệu thuần (Appointment, Patient, Invoice, MedicalRecord,...)',
    'Filter: AuthenticationFilter (xác thực) → AuthorizationFilter (phân quyền) → CsrfFilter (chống CSRF) → EncodingFilter (UTF-8)',
    'Utility: AuditUtil (ghi audit log), NotificationHelper (tạo thông báo), BCryptUtil (mã hoá), EncryptionUtil (mã hoá dữ liệu nhạy cảm)',
]
for item in arch_items:
    add_bullet(item)

add_heading('2.2. Cấu trúc thư mục', 2)
add_table(
    ['Thư mục', 'Mô tả'],
    [
        ['src/java/com/clinic/controller/', '60+ Servlet xử lý HTTP request'],
        ['src/java/com/clinic/service/', '21 Service class chứa logic nghiệp vụ'],
        ['src/java/com/clinic/dao/', '33 DAO class truy vấn database'],
        ['src/java/com/clinic/model/', '35+ Model class (POJO + enums)'],
        ['src/java/com/clinic/filter/', '4 Filter (Auth, Authorization, Csrf, Encoding)'],
        ['src/java/com/clinic/config/', 'Cấu hình DB, Google OAuth, SlotHold listener'],
        ['src/java/com/clinic/utils/', 'Tiện ích: Audit, Notification, BCrypt, Encryption, etc.'],
        ['web/views/', '65+ file JSP chia theo actor (admin, manager, doctors, staff, sonographer, patient, auth, common, errors, home)'],
        ['web/WEB-INF/', 'web.xml, cấu hình deployment'],
        ['docs/', 'Đặc tả use case, prompt triển khai, schema database, kịch bản test'],
    ],
    [5, 12]
)

add_heading('2.3. Công nghệ sử dụng', 2)
add_table(
    ['Tầng', 'Công nghệ'],
    [
        ['Backend', 'Java 17+, Jakarta Servlet API, JSP, JSTL'],
        ['Database', 'SQL Server (Microsoft SQL Server)'],
        ['Frontend', 'JSP, HTML5, CSS3, JavaScript, Bootstrap'],
        ['Authentication', 'Session-based, Google OAuth 2.0, BCrypt password hashing'],
        ['AI Engine', 'Python AI service (Flask) — phân tích ảnh siêu âm qua HTTP'],
        ['Container', 'Apache Tomcat 10+'],
        ['Build', 'Ant / IntelliJ IDEA'],
        ['Version Control', 'Git (GitHub)'],
    ],
    [4, 13]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN III: CƠ SỞ DỮ LIỆU
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN III: CƠ SỞ DỮ LIỆU', 1)

add_heading('3.1. Danh sách bảng', 2)
add_table(
    ['#', 'Tên bảng', 'Mô tả', 'Số cột'],
    [
        ['1', 'users', 'Tài khoản người dùng (email, password, role, provider, verified)', '17'],
        ['2', 'roles', 'Danh sách vai trò (Admin, Manager, Doctor, Staff, Sonographer, Patient)', '3'],
        ['3', 'permissions', 'Danh sách quyền chi tiết (RBAC)', '6'],
        ['4', 'role_permissions', 'Ánh xạ Role ↔ Permission', '3'],
        ['5', 'patients', 'Hồ sơ bệnh nhân (tên, SĐT, ngày sinh, CCCD, địa chỉ)', '8'],
        ['6', 'doctors', 'Hồ sơ bác sĩ (chuyên khoa, bằng cấp, năm kinh nghiệm)', '8'],
        ['7', 'sonographers', 'Hồ sơ bác sĩ siêu âm (kinh nghiệm, chứng chỉ, phòng)', '5'],
        ['8', 'appointments', 'Lịch hẹn khám (core table) — 22 cột', '22'],
        ['9', 'appointment_services', 'Dịch vụ đi kèm lịch hẹn (service_id, price)', '4'],
        ['10', 'medical_records', 'Hồ sơ bệnh án (sinh hiệu, khám lâm sàng, chẩn đoán)', '31'],
        ['11', 'test_orders', 'Chỉ định cận lâm sàng (siêu âm) — service_id, lý do', '10'],
        ['12', 'prescriptions', 'Đơn thuốc (mã, trạng thái, quyết định mua)', '7'],
        ['13', 'prescription_items', 'Chi tiết đơn thuốc (thuốc, số lượng, liều)', '5'],
        ['14', 'medicines', 'Danh mục thuốc (tên, giá, tồn kho, mã)', '12'],
        ['15', 'medicine_categories', 'Danh mục nhóm thuốc', '7'],
        ['16', 'medicine_price_history', 'Lịch sử đổi giá thuốc', '6'],
        ['17', 'invoices', 'Hoá đơn (PRE_EXAM, POST_EXAM, PRESCRIPTION) — 12 cột', '12'],
        ['18', 'invoice_items', 'Chi tiết hoá đơn (item_type, quantity, unit_price)', '7'],
        ['19', 'services', 'Danh mục dịch vụ y tế (tên, giá, mã, yêu cầu)', '13'],
        ['20', 'service_categories', 'Danh mục nhóm dịch vụ', '7'],
        ['21', 'price_history', 'Lịch sử đổi giá dịch vụ', '6'],
        ['22', 'doctor_schedules', 'Lịch làm việc bác sĩ (ca, ngày, booked_count)', '19'],
        ['23', 'shifts', 'Ca làm việc (tên, giờ bắt đầu/kết thúc)', '7'],
        ['24', 'ultrasound_images', 'Ảnh siêu âm đã tải lên', '10'],
        ['25', 'ultrasound_annotations', 'Chú thích/vẽ thủ công trên ảnh siêu âm', '16'],
        ['26', 'ultrasound_reports', 'Báo cáo kết quả siêu âm (mô tả, kết luận, chữ ký)', '15'],
        ['27', 'ai_analysis_results', 'Kết quả phân tích AI (bounding box, confidence)', '16'],
        ['28', 'audit_logs', 'Nhật ký hệ thống (user, action, old/new value, IP)', '7'],
        ['29', 'notifications', 'Thông báo (user, title, content, channel, read)', '6'],
        ['30', 'pregnancies', 'Thai kỳ (ngày bắt đầu, dự sinh, trạng thái)', '8'],
        ['31', 'reviews', 'Đánh giá của bệnh nhân (rating, comment)', '4'],
        ['32', 'password_reset_tokens', 'Token đặt lại mật khẩu', '6'],
    ],
    [1, 4, 8, 1.5]
)

add_heading('3.2. Chi tiết các bảng chính', 2)

add_para('Bảng appointments (22 cột) — Bảng trung tâm của hệ thống:', bold=True)
add_table(
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['id', 'int IDENTITY', 'PK — Mã lịch hẹn'],
        ['patient_id', 'int', 'FK → patients.id'],
        ['doctor_id', 'int', 'FK → doctors.id'],
        ['pregnancy_id', 'int', 'FK → pregnancies.id (nếu có thai kỳ)'],
        ['appointment_date', 'date', 'Ngày khám'],
        ['booking_source', 'nvarchar(50)', '"Online" hoặc "Manual"'],
        ['symptoms', 'nvarchar(max)', 'Triệu chứng/lý do khám (10-500 ký tự)'],
        ['last_menstrual_period', 'date', 'Ngày kinh cuối (LMP) — cho bệnh nhân nữ'],
        ['is_emergency', 'bit', 'Đánh dấu ca cấp cứu'],
        ['status', 'nvarchar(30)', 'Pending, Confirmed, Waiting, InProgress, SUCCESS, Completed, Cancelled, NoShow'],
        ['service_id', 'int', 'FK → services.id (dịch vụ khám chính)'],
        ['time_slot', 'time(7)', 'Giờ khám'],
        ['queue_number', 'nvarchar(20)', 'Số thứ tự hàng đợi'],
        ['base_fee', 'decimal(12,2)', 'GIÁ KHÁM KHOÁ CỨNG — lưu tại thời điểm đặt lịch, không đổi khi Manager sửa bảng giá'],
        ['priority_reason', 'nvarchar(500)', 'Lý do đánh dấu ưu tiên (5-500 ký tự)'],
        ['prioritized_at', 'datetime2', 'Thời điểm đánh dấu ưu tiên'],
        ['prioritized_by', 'int', 'FK → users.id — ai đánh dấu ưu tiên'],
        ['is_priority', 'bit', 'Cờ ưu tiên'],
        ['schedule_id', 'int', 'FK → doctor_schedules.id'],
        ['held_by', 'int', 'FK → users.id — ai đang giữ slot (chống race condition)'],
        ['held_until', 'datetime2', 'Slot được giữ đến khi nào'],
    ],
    [4, 3, 10]
)

add_para('')
add_para('Bảng invoices (12 cột):', bold=True)
add_table(
    ['Cột', 'Kiểu', 'Mô tả'],
    [
        ['id', 'int IDENTITY', 'PK'],
        ['appointment_id', 'int', 'FK → appointments.id'],
        ['total_amount', 'decimal(18,2)', 'Tổng tiền'],
        ['status', 'nvarchar(30)', 'Unpaid, Paid, Refunded, Cancelled'],
        ['transaction_code', 'nvarchar(100)', 'Mã giao dịch'],
        ['invoice_type', 'varchar(30)', 'PRE_EXAM (phí khám), POST_EXAM (dịch vụ CLS), PRESCRIPTION (thuốc)'],
        ['payment_method', 'nvarchar(30)', 'CASH (mặc định, duy nhất)'],
        ['confirmed_by', 'int', 'FK → users.id — staff xác nhận thanh toán'],
        ['confirmed_at', 'datetime', 'Thời điểm xác nhận'],
        ['paid_at', 'datetime', 'Thời điểm thanh toán — bằng chứng thu tiền'],
        ['paid_by_user_id', 'int', 'FK → users.id — staff thu tiền'],
        ['payment_note', 'nvarchar(500)', 'Ghi chú thanh toán'],
    ],
    [3.5, 3, 10.5]
)

add_para('')
add_para('Bảng doctor_schedules (19 cột) — thay thế time_slots cũ:', bold=True)
add_table(
    ['Cột', 'Mô tả'],
    [
        ['id', 'PK'],
        ['doctor_id', 'FK → doctors.id'],
        ['work_date', 'Ngày làm việc'],
        ['shift_id', 'FK → shifts.id — ca làm việc'],
        ['max_slots', 'Số bệnh nhân tối đa trong ca'],
        ['booked_count', 'Số bệnh nhân đã đặt — tăng/giảm atomic trong transaction'],
        ['is_approved', 'Manager đã duyệt chưa'],
        ['status', 'DRAFT, PENDING, APPROVED, REJECTED, CANCELLED'],
        ['rejection_reason', 'Lý do từ chối'],
        ['approved_by / approved_at', 'Người duyệt & thời điểm'],
        ['created_by / created_at', 'Người tạo & thời điểm'],
        ['cancelled_by / cancelled_at / cancellation_reason', 'Thông tin huỷ ca'],
        ['version', 'timestamp — optimistic locking'],
        ['notes', 'Ghi chú'],
    ],
    [4.5, 12.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN IV: USE CASE CHI TIẾT
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN IV: USE CASE CHI TIẾT', 1)

add_heading('4.1. Use Case Diagram (tổng quan)', 2)
add_para('Hệ thống có 6 Actor chính với các use case sau (đã chốt theo đặc tả cams-dac-ta-use-case.md):')

# ── 4.2 PATIENT ──
add_heading('4.2. Patient (Bệnh nhân)', 2)
add_table(
    ['#', 'Use Case', 'Loại', 'Mô tả', 'File Controller chính'],
    [
        ['UC-P1', 'Book Appointment', 'Core', 'Đặt lịch khám online: chọn bác sĩ → xem giá → chọn slot → nhập triệu chứng/LMP → xác nhận', 'PatientBookingServlet, PatientSlotApiServlet'],
        ['UC-P2', 'View My Appointments', 'Core', 'Xem danh sách lịch hẹn, trạng thái (tiếng Việt), số thứ tự hàng đợi, khối việc cần làm', 'PatientAppointmentServlet'],
        ['UC-P3', 'Cancel Appointment', 'Core', 'Huỷ lịch hẹn: điều kiện Pending/Confirmed, chưa thanh toán, trước 2h, chưa check-in', 'PatientAppointmentServlet'],
        ['UC-P4', 'Change My Appointment', 'Core', 'Đổi slot cùng bác sĩ: tối đa 2 lần, giữ nguyên giá, trả slot cũ + nhận slot mới', 'PatientAppointmentServlet'],
        ['UC-P5', 'View Medical Record', 'Core', 'Xem hồ sơ bệnh án đã chốt (Final), chỉ đọc, kiểm tra chủ sở hữu', 'PatientMedicalRecordServlet'],
        ['UC-P6', 'View Payment History', 'Core', 'Xem lịch sử thanh toán, hoá đơn', 'PatientInvoicesServlet, PatientPaymentServlet'],
        ['UC-P7', 'Manage Profile', 'Support', 'Cập nhật thông tin cá nhân', 'PatientProfileServlet'],
        ['UC-P8', 'Review Doctor', 'Support', 'Đánh giá bác sĩ sau khám (rating + comment)', 'PatientReviewServlet'],
    ],
    [0.5, 3, 1.2, 7, 5]
)

# Use Case Specs for Patient
add_heading('Đặc tả UC-P1: Book Appointment', 3)
add_para('Tiền điều kiện:', bold=True)
add_bullet('Đã đăng nhập vai trò Patient')
add_bullet('Hồ sơ cá nhân có đủ họ tên, số điện thoại, giới tính')
add_para('Luồng chính:', bold=True)
add_bullet('1. Bệnh nhân chọn bác sĩ → hiển thị giá khám của bác sĩ đó (dựa trên experience_years)')
add_bullet('2. Chọn ngày → hiển thị các slot còn chỗ của lịch làm việc đã được Manager duyệt')
add_bullet('3. Chọn slot → nhập triệu chứng (bắt buộc 10-500 ký tự), nhập LMP nếu là bệnh nhân nữ')
add_bullet('4. Xác nhận: hiển thị toàn bộ thông tin + số tiền + dòng "thanh toán tại quầy lễ tân trước giờ khám"')
add_bullet('5. Tạo lịch hẹn Pending, giá khám (base_fee) được lưu cứng, booked_count tăng 1, audit log')
add_para('Luồng thay thế:', bold=True)
add_bullet('Slot vừa hết chỗ → báo lỗi, tải lại danh sách')
add_bullet('Đã có lịch cùng ngày còn hiệu lực → từ chối (1 ngày/1 bệnh nhân/1 lịch)')
add_bullet('Slot quá khứ hoặc sắp bắt đầu → từ chối')
add_bullet('Bấm 2 lần/F5 → chỉ tạo 1 lịch (idempotency)')
add_para('Hậu điều kiện:', bold=True)
add_bullet('Lịch hẹn Pending, base_fee đã lưu cứng, chưa có hoá đơn, bác sĩ nhận thông báo')

add_heading('Đặc tả UC-P3: Cancel Appointment (Patient)', 3)
add_para('Tiền điều kiện (KIỂM TRA Ở SERVER):', bold=True)
add_bullet('Lịch hẹn thuộc về chính bệnh nhân đang đăng nhập')
add_bullet('Trạng thái Pending hoặc Confirmed')
add_bullet('Hoá đơn PRE_EXAM CHƯA thanh toán')
add_bullet('Chưa check-in')
add_bullet('Còn cách giờ bắt đầu ca ít nhất 2 giờ')
add_para('Luồng chính:', bold=True)
add_bullet('Bấm Huỷ → hộp xác nhận (không thể hoàn tác) → nhập lý do (không bắt buộc, max 500 ký tự) → xác nhận')
add_para('Hậu điều kiện (transaction nguyên tử):', bold=True)
add_bullet('Trạng thái → Cancelled, hoá đơn chưa thanh toán bị huỷ, booked_count giảm 1, audit log, thông báo bác sĩ')

add_heading('Đặc tả UC-P4: Change My Appointment (Patient)', 3)
add_para('Điều kiện: giống UC-P3 (huỷ lịch). Giới hạn: CHỈ đổi slot cùng bác sĩ, tối đa 2 lần (đếm từ audit log).')
add_para('Transaction:', bold=True)
add_bullet('Giảm booked_count slot cũ + tăng booked_count slot mới + cập nhật lịch hẹn')
add_bullet('Giá khám GIỮ NGUYÊN (cùng bác sĩ), hoá đơn PRE_EXAM giữ nguyên')
add_bullet('Nếu slot mới vừa hết → rollback toàn bộ')

doc.add_page_break()

# ── 4.3 STAFF ──
add_heading('4.3. Staff (Nhân viên lễ tân)', 2)
add_table(
    ['#', 'Use Case', 'Mô tả', 'Điều kiện chính'],
    [
        ['UC-S1', 'View Reception Queue', 'Xem hàng đợi lễ tân: lọc theo ngày, trạng thái, tìm kiếm; sắp xếp ưu tiên', '—'],
        ['UC-S2', 'Approve Booking', 'Duyệt lịh hẹn: Pending → Confirmed + Tạo hoá đơn PRE_EXAM chưa thanh toán (copy base_fee)', 'Trạng thái Pending'],
        ['UC-S3', 'Confirm Payment (PRE_EXAM)', 'Xác nhận thu tiền mặt phí khám: ghi paid_at, paid_by_user_id, payment_method=CASH', 'Hoá đơn Unpaid; có PRE_EXAM'],
        ['UC-S4', 'Check In Patient', 'Check-in: Confirmed → Waiting + cấp số thứ tự + xếp hàng đợi', 'Đã duyệt, hôm nay, PRE_EXAM ĐÃ thanh toán, chưa check-in, trong khung giờ'],
        ['UC-S5', 'Mark Priority', 'Đánh dấu ưu tiên: nhập lý do (5-500 ký tự) + xếp lại hàng đợi', 'Waiting, hôm nay, đã thanh toán, chưa ưu tiên'],
        ['UC-S6', 'Cancel Appointment', 'Huỷ lịch: nếu PRE_EXAM đã thanh toán thì hoàn tiền + huỷ', 'Chưa InProgress, chưa SUCCESS'],
        ['UC-S7', 'Mark No Show', 'Đánh dấu không đến: qua giờ kết thúc ca hôm nay', 'Pending/Confirmed, hôm nay, ca đã kết thúc'],
        ['UC-S8', 'Create Manual Booking', 'Tạo lịch thủ công cho bệnh nhân vãng lai', '—'],
        ['UC-S9', 'Update Appointment Info', 'Sửa thông tin lịch hẹn (chỉ Pending, chưa có hoá đơn)', 'Pending + PRE_EXAM Unpaid'],
        ['UC-S10', 'Refund Invoice', 'Hoàn tiền hoá đơn đã thanh toán: Paid → Refunded + lý do', 'Hoá đơn Paid, KHÔNG hoàn 2 lần'],
        ['UC-S11', 'Create Initial Patient Record', 'Tạo hồ sơ hành chính cho bệnh nhân mới (KHÔNG được nhập chẩn đoán)', '—'],
        ['UC-S12', 'Confirm POST_EXAM Payment', 'Xác nhận thanh toán dịch vụ siêu âm', 'InProgress, POST_EXAM Unpaid'],
        ['UC-S13', 'View Doctor Schedule', 'Xem lịch làm việc bác sĩ (read-only)', '—'],
    ],
    [0.5, 3.5, 8, 5]
)

add_heading('Đặc tả UC-S2: Approve Booking', 3)
add_para('Tiền điều kiện: Lịch hẹn Pending')
add_para('Luồng chính:')
add_bullet('1. Staff xem hàng đợi → bấm "Duyệt & Tạo Hoá Đơn"')
add_bullet('2. Hệ thống chuyển sang Confirmed + tạo hoá đơn PRE_EXAM Unpaid, số tiền copy từ base_fee')
add_bullet('3. Nếu đã có PRE_EXAM → không tạo trùng')
add_para('Hậu điều kiện: Confirmed, tồn tại đúng 1 PRE_EXAM Unpaid, bệnh nhân nhận thông báo')

add_heading('Đặc tả UC-S4: Check In Patient', 3)
add_para('Tiền điều kiện BẮT BUỘC ĐỦ (KIỂM TRA Ở SERVER):')
add_bullet('Lịch hẹn Confirmed')
add_bullet('Của NGÀY HÔM NAY')
add_bullet('Hoá đơn PRE_EXAM ĐÃ thanh toán (Paid) — nếu chưa: "Bệnh nhân chưa thanh toán phí khám"')
add_bullet('Chưa check-in (không phải Waiting)')
add_bullet('Trong khung giờ cho phép check-in (tối đa 120 phút trước giờ ca)')
add_bullet('Ca chưa kết thúc')
add_para('Hậu điều kiện: Waiting + số thứ tự, xếp hàng đợi theo ưu tiên → giờ, bệnh nhân xem được số')

add_heading('Đặc tả UC-S10: Refund Invoice', 3)
add_para('Điều kiện: Hoá đơn Paid. Không hoàn 2 lần.')
add_para('Hậu quả:')
add_bullet('Hoá đơn → Refunded, ghi refunded_at, refunded_by_user_id, refund_reason')
add_bullet('Nếu là POST_EXAM → tự động huỷ test_order liên quan')
add_bullet('Doanh thu Manager bị trừ khoản này')
add_bullet('Audit log + thông báo cho bệnh nhân')

doc.add_page_break()

# ── 4.4 DOCTOR ──
add_heading('4.4. Doctor (Bác sĩ lâm sàng)', 2)
add_table(
    ['#', 'Use Case', 'Mô tả', 'Giai đoạn (ExamStage)'],
    [
        ['UC-D1', 'View Today\'s Appointments', 'Xem danh sách bệnh nhân hôm nay + tiếp nhận ca (Waiting → InProgress)', 'NOT_STARTED'],
        ['UC-D2', 'Update Medical Record', 'Khám lâm sàng: sinh hiệu, triệu chứng, tiền sử → lưu nháp', 'CLINICAL_EXAM / ORDER_DECISION'],
        ['UC-D3', 'Order Ultrasound', 'Chỉ định siêu âm: chọn dịch vụ + lý do (10-500 ký tự) → tạo test_order + POST_EXAM', 'ORDER_DECISION'],
        ['UC-D4', 'Cancel Ultrasound Order', 'Huỷ chỉ định: chỉ khi POST_EXAM chưa thanh toán + sonographer chưa nhận ca', 'WAITING_PAYMENT'],
        ['UC-D5', 'Diagnose and Prescribe', 'Chẩn đoán + kê đơn: xem kết quả siêu âm (read-only) → nhập chẩn đoán → kê đơn', 'DIAGNOSIS / READY_TO_FINALIZE'],
        ['UC-D6', 'Finalize Medical Record', 'Chốt hồ sơ: medical_record → Final, appointment → SUCCESS, xoá STT', 'READY_TO_FINALIZE'],
        ['UC-D7', 'View Medical Record Details', 'Xem chi tiết hồ sơ (đã chốt thì read-only)', '—'],
        ['UC-D8', 'Manage Work Schedules', 'Đăng ký/huỷ lịch làm việc cá nhân', '—'],
        ['UC-D9', 'View Appointment Overview', 'Dashboard tổng quan', '—'],
        ['UC-D10', 'Manage Prescriptions', 'Xem danh sách đơn thuốc đã kê', '—'],
    ],
    [0.5, 3.5, 8, 5]
)

add_heading('Đặc tả UC-D2 → UC-D6: Luồng khám 5 khối (theo getStage)', 3)
add_para('QUY TẮC HIỂN THỊ:')
add_bullet('Khối đã xong: thu gọn, dấu tích, tóm tắt, có nút "Sửa"')
add_bullet('Khối đang tới lượt: mở rộng, nhập được')
add_bullet('Khối chưa tới lượt: thu gọn, xám, biểu tượng khoá, GHI RÕ LÝ DO bằng tiếng Việt')
add_bullet('Cho phép quay lại SỬA khối trước cho tới khi chốt. Sau chốt → toàn bộ read-only')
add_para('')
add_para('KHỐI 1 — Tiếp nhận ca:', bold=True)
add_bullet('Chỉ có nút "Tiếp nhận", không form. Waiting → InProgress.')
add_para('KHỐI 2 — Khám lâm sàng:', bold=True)
add_bullet('Form: triệu chứng, tiền sử, sinh hiệu (mạch, HA, nhiệt độ, cân nặng, tuổi thai, ...). Nút "Lưu nháp".')
add_bullet('Sau khi lưu → mở khoá Khối 3.')
add_para('KHỐI 3 — Chỉ định cận lâm sàng:', bold=True)
add_bullet('2 lựa chọn: "Không chỉ định" hoặc "Chỉ định siêu âm"')
add_bullet('Nếu chỉ định: chọn dịch vụ từ danh mục + lý do BẮT BUỘC 10-500 ký tự')
add_bullet('Sau khi chỉ định → khoá Khối 4, 5. Hiện lý do: "Đang chờ bệnh nhân thanh toán" hoặc "Đang chờ kết quả siêu âm"')
add_bullet('Có nút "Huỷ chỉ định" (chỉ khi POST_EXAM chưa thanh toán)')
add_para('KHỐI 4 — Chẩn đoán và kê đơn:', bold=True)
add_bullet('Bên trên: kết quả siêu âm CHỈ ĐỌC (ảnh, AI, kết luận sonographer). BS lâm sàng KHÔNG sửa được.')
add_bullet('Bên dưới: chẩn đoán, kết luận, kê đơn, lựa chọn mua tại PK hay mua ngoài')
add_bullet('Kê đơn KHÔNG chặn chốt hồ sơ. Hoá đơn thuốc do staff xử lý riêng.')
add_para('KHỐI 5 — Chốt hồ sơ:', bold=True)
add_bullet('Nút "Chốt hồ sơ bệnh án" — transaction nguyên tử: medical_record → Final, appointment → SUCCESS, xoá STT')

doc.add_page_break()

# ── 4.5 SONOGRAPHER ──
add_heading('4.5. Sonographer (Bác sĩ siêu âm)', 2)
add_table(
    ['#', 'Use Case', 'Mô tả', 'Thứ tự'],
    [
        ['UC-SN1', 'View Ultrasound Overview', 'Dashboard danh sách ca siêu âm đang chờ/đang xử lý', '—'],
        ['UC-SN2', 'Accept Ultrasound Case', 'Tiếp nhận ca: kiểm tra isReadyForSonographer → start', 'Bước 1'],
        ['UC-SN3', 'Capture and Upload Images', 'Tải ảnh siêu âm lên (lưu bytes gốc, KHÔNG re-encode)', 'Bước 2'],
        ['UC-SN4', 'Analyze with AI', 'Gửi ảnh sang AI Engine (HTTP) → phân tích → kết quả (bounding box, confidence)', 'Bước 3'],
        ['UC-SN5', 'Accept AI Result', 'Đồng ý kết quả AI → điền thông tin → ký', 'Bước 3a'],
        ['UC-SN6', 'Reject AI Result → Draw Manually', 'Từ chối AI → BẮT BUỘC vẽ thủ công (annotation) → điền thông tin → ký', 'Bước 3b'],
        ['UC-SN7', 'Skip AI (AI lỗi)', 'AI Engine lỗi/timeout → vẽ thủ công hoàn toàn, ghi log', 'Bước 3 (alt)'],
        ['UC-SN8', 'Update Ultrasound Result', 'Điền mô tả, phát hiện chuyên môn, kết luận', 'Bước 3/4'],
        ['UC-SN9', 'Sign and Confirm', 'Ký xác nhận: kiểm tra đã có vẽ (nếu từ chối AI), đã có thông tin', 'Bước 4'],
        ['UC-SN10', 'Publish Results', 'Công bố: test_order → Completed + lưu report + THÔNG BÁO bác sĩ lâm sàng', 'Bước 4'],
        ['UC-SN11', 'View Medical Record Details', 'Xem hồ sơ bệnh án liên quan (read-only)', '—'],
    ],
    [0.5, 3.5, 8, 1.2]
)

add_heading('Đặc tả luồng siêu âm (4 bước tuần tự, không bỏ bước)', 3)
add_para('Bước 1 — Tiếp nhận ca:', bold=True)
add_bullet('Kiểm tra isReadyForSonographer (POST_EXAM đã thanh toán), chưa có sonographer nào nhận')
add_bullet('Gán sonographer_user_id → trạng thái "InProgress"')
add_para('Bước 2 — Tải ảnh:', bold=True)
add_bullet('Upload ảnh siêu âm, lưu bytes gốc (KHÔNG re-encode qua ImageIO), chỉ avatar mới re-encode')
add_para('Bước 3 — Phân tích AI:', bold=True)
add_bullet('Gửi HTTP request tới AI Engine (Python Flask), có timeout cụ thể')
add_bullet('Nếu AI lỗi/timeout: cho phép bỏ qua, vẽ thủ công hoàn toàn, ghi log "không qua AI"')
add_bullet('2 nhánh: Accept AI Result (đồng ý) hoặc Reject AI Result (từ chối → BẮT BUỘC vẽ)')
add_bullet('Validate SERVER: chọn vẽ tay thì PHẢI có ít nhất 1 annotation mới cho ký')
add_para('Bước 4 — Ký & Công bố:', bold=True)
add_bullet('Điền image_description, professional_findings, conclusion → ký (signed_by_user_id, signed_name, signed_at)')
add_bullet('Công bố: test_order → Completed + lưu ultrasound_report + GỬI THÔNG BÁO CHO BÁC SĨ LÂM SÀNG')
add_bullet('Đây là mắt xích bàn giao quan trọng nhất giữa hai bác sĩ')

doc.add_page_break()

# ── 4.6 MANAGER ──
add_heading('4.6. Manager (Quản lý)', 2)
add_table(
    ['#', 'Use Case', 'Mô tả'],
    [
        ['UC-M1', 'Manage Medical Services', 'Thêm/sửa/tắt dịch vụ y tế. Đổi giá → ghi price_history. KHÔNG xoá cứng. Giá mới KHÔNG ảnh hưởng hoá đơn cũ.'],
        ['UC-M2', 'View Price Adjustment History', 'Xem lịch sử đổi giá: lọc theo dịch vụ + khoảng thời gian, hiển thị giá cũ/mới, chênh lệch, người đổi.'],
        ['UC-M3', 'Manage Revenue', 'Xem doanh thu: đọc TRỰC TIẾP từ invoices (Paid - Refunded), lọc khoảng ngày, tách 3 loại PRE_EXAM/POST_EXAM/PRESCRIPTION.'],
        ['UC-M4', 'View Service Details', 'Chi tiết từng hoá đơn: bệnh nhân, dịch vụ, số tiền, thời điểm, staff xác nhận.'],
        ['UC-M5', 'View Service Statistics', 'Thống kê SỐ LƯỢT (không phải tiền): số lượt chỉ định + thực hiện theo dịch vụ, theo thời gian.'],
        ['UC-M6', 'Manage Work Schedules', 'Duyệt/từ chối lịch làm việc của bác sĩ, xem lịch theo ngày.'],
        ['UC-M7', 'Manage Appointment Time Slots', 'Quản lý khung giờ/doctor_schedules.'],
        ['UC-M8', 'Manage Doctors', 'Xem danh sách + chi tiết bác sĩ.'],
        ['UC-M9', 'Manage Medicines', 'Quản lý danh mục thuốc (giống Admin).'],
        ['UC-M10', 'Monitor Stuck Appointments', 'Danh sách ca chưa chốt quá 24h: bệnh nhân, bác sĩ, thời điểm, đang kẹt ở bước nào (dùng getStage).'],
    ],
    [0.5, 3.5, 13]
)

add_para('')
add_para('Lưu ý quan trọng về Revenue vs Statistics:', bold=True)
add_bullet('Revenue (UC-M3) = TIỀN đã thu (Paid) - TIỀN đã hoàn (Refunded), đọc trực tiếp từ invoices, KHÔNG tạo bảng tổng hợp')
add_bullet('Statistics (UC-M5) = SỐ LƯỢT dịch vụ được chỉ định và thực hiện, theo khoảng thời gian')
add_bullet('Hai màn hình TÁCH BIỆT, không trộn vào một')

# ── 4.7 ADMIN ──
add_heading('4.7. Admin (Quản trị viên)', 2)
add_table(
    ['#', 'Use Case', 'Mô tả'],
    [
        ['UC-A1', 'Manage Users', 'Thêm/sửa/khoá/xoá mềm tài khoản, gán vai trò'],
        ['UC-A2', 'Manage Roles & Permissions', 'Quản lý vai trò, phân quyền chi tiết (RBAC)'],
        ['UC-A3', 'Manage Medicines', 'Quản lý danh mục thuốc + giá + tồn kho'],
        ['UC-A4', 'Manage Services', 'Quản lý danh mục dịch vụ y tế'],
        ['UC-A5', 'View Audit Logs', 'Xem nhật ký hệ thống: lọc theo user, action, thời gian'],
        ['UC-A6', 'Manage Pricing', 'Quản lý giá (pricing configuration)'],
        ['UC-A7', 'System Dashboard', 'Dashboard tổng quan hệ thống'],
    ],
    [0.5, 3.5, 13]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN V: LUỒNG XỬ LÝ NGHIỆP VỤ
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN V: LUỒNG XỬ LÝ NGHIỆP VỤ', 1)

add_heading('5.1. Luồng chính: Đặt lịch → Hoàn tất khám', 2)
add_para('Đây là luồng end-to-end xuyên suốt toàn bộ hệ thống, qua 5 actor:')
add_para('')
steps = [
    ('1. ĐẶT LỊCH (Patient)', 'Patient chọn bác sĩ → xem giá → chọn slot → nhập triệu chứng/LMP → Xác nhận.\n'
     'KẾT QUẢ: appointment Pending, base_fee lưu cứng, booked_count +1, chưa có invoice.\n'
     'VALIDATE: slot còn chỗ, không trùng lịch trong ngày, triệu chứng 10-500 ký tự, LMP hợp lệ.'),
    ('2. DUYỆT LỊCH (Staff)', 'Staff bấm "Duyệt & Tạo Hoá Đơn".\n'
     'KẾT QUẢ: appointment Confirmed, invoice PRE_EXAM Unpaid (copy base_fee).\n'
     'VALIDATE: appointment Pending, chưa có PRE_EXAM (chống trùng).'),
    ('3. THU TIỀN PHÍ KHÁM (Staff)', 'Staff bấm "Xác nhận thanh toán".\n'
     'KẾT QUẢ: invoice PRE_EXAM Paid, ghi paid_at, paid_by_user_id, payment_method=CASH.\n'
     'VALIDATE: invoice Unpaid, chưa Paid (chống xác nhận trùng), staff quyền Staff.'),
    ('4. CHECK-IN (Staff)', 'Staff bấm "Check-in".\n'
     'KẾT QUẢ: appointment Waiting, cấp queue_number, xếp hàng đợi.\n'
     'VALIDATE: Confirmed, hôm nay, PRE_EXAM Paid, chưa Waiting, trong khung giờ.'),
    ('5. TIẾP NHẬN CA (Doctor)', 'Doctor bấm "Tiếp nhận" từ danh sách bệnh nhân hôm nay.\n'
     'KẾT QUẢ: appointment InProgress.\n'
     'VALIDATE: Waiting, doctor được phân công.'),
    ('6. KHÁM LÂM SÀNG (Doctor)', 'Doctor nhập sinh hiệu, triệu chứng, tiền sử → "Lưu nháp".\n'
     'KẾT QUẢ: medical_record Draft.\n'
     'VALIDATE: InProgress, đúng doctor. Mở khoá Khối 3.'),
    ('7a. KHÔNG CHỈ ĐỊNH (Doctor)', 'Doctor chọn "Không chỉ định" → chuyển thẳng sang Khối 4 (Chẩn đoán).\n'
     'VALIDATE: ORDER_DECISION stage.'),
    ('7b. CHỈ ĐỊNH SIÊU ÂM (Doctor)', 'Doctor chọn dịch vụ + nhập lý do (10-500 ký tự).\n'
     'KẾT QUẢ: test_order Pending, invoice POST_EXAM Unpaid.\n'
     'VALIDATE: ORDER_DECISION stage. Khoá Khối 4,5.'),
    ('8. THU TIỀN SIÊU ÂM (Staff)', 'Staff xác nhận thanh toán POST_EXAM.\n'
     'KẾT QUẢ: POST_EXAM Paid. Mở WAITING_ULTRASOUND.\n'
     'VALIDATE: POST_EXAM Unpaid.'),
    ('9. SIÊU ÂM (Sonographer)', 'Sonographer: nhận ca → tải ảnh → AI → ký → công bố.\n'
     'KẾT QUẢ: test_order Completed, có report, THÔNG BÁO doctor. Mở khoá Khối 4.\n'
     'VALIDATE: đúng sonographer, tuần tự 4 bước, có annotation nếu vẽ tay.'),
    ('10. CHẨN ĐOÁN & KÊ ĐƠN (Doctor)', 'Doctor xem kết quả siêu âm (read-only) → nhập chẩn đoán → kê đơn.\n'
     'KẾT QUẢ: final_diagnosis, prescription (nếu có), purchase_decision.\n'
     'VALIDATE: DIAGNOSIS/READY_TO_FINALIZE stage. Mở khoá Khối 5.'),
    ('11. CHỐT HỒ SƠ (Doctor)', 'Doctor bấm "Chốt hồ sơ".\n'
     'KẾT QUẢ (transaction): medical_record Final, appointment SUCCESS, xoá queue_number.\n'
     'VALIDATE: READY_TO_FINALIZE stage. KHÔNG yêu cầu thanh toán thuốc.'),
]
for title, detail in steps:
    add_para(title, bold=True)
    add_para(detail)

doc.add_page_break()

add_heading('5.2. Sơ đồ chuyển đổi trạng thái lịch hẹn', 2)
add_para('Trạng thái lịch hẹn (appointment.status) và điều kiện chuyển đổi:')
add_para('')
add_table(
    ['Từ trạng thái', 'Đến trạng thái', 'Actor', 'Điều kiện', 'Side effects'],
    [
        ['(mới tạo)', 'Pending', 'Patient/Staff', 'Slot còn chỗ, không trùng lịch', 'booked_count +1, base_fee lưu cứng'],
        ['Pending', 'Confirmed', 'Staff', 'Staff duyệt (Approve Booking)', 'Tạo PRE_EXAM Unpaid (copy base_fee)'],
        ['Confirmed', 'Waiting', 'Staff', 'Hôm nay + PRE_EXAM Paid + trong khung giờ', 'Cấp STT, xếp hàng đợi'],
        ['Waiting', 'InProgress', 'Doctor', 'Doctor được phân công', 'Bắt đầu khám'],
        ['InProgress', 'SUCCESS', 'Doctor', 'Chốt hồ sơ (Finalize)', 'MR Final, xoá STT'],
        ['Pending', 'Cancelled', 'Patient/Staff', 'PRE_EXAM chưa thanh toán, trước 2h (patient)', 'booked_count -1, huỷ invoice Unpaid'],
        ['Confirmed', 'Cancelled', 'Patient/Staff', 'PRE_EXAM chưa thanh toán, trước 2h (patient)', 'booked_count -1, huỷ invoice; nếu Paid→Refund'],
        ['Pending/Confirmed', 'NoShow', 'Staff', 'Hôm nay, ca đã kết thúc', 'booked_count -1, huỷ invoice Unpaid; nếu Paid→cảnh báo'],
    ],
    [2.5, 2.5, 1.5, 5, 5.5]
)

add_heading('5.3. Luồng thanh toán', 2)
add_para('Hệ thống có 3 loại hoá đơn (invoice_type):', bold=True)
add_bullet('PRE_EXAM: Phí khám — tạo khi staff duyệt lịch, số tiền copy từ base_fee (không tính lại)')
add_bullet('POST_EXAM: Dịch vụ cận lâm sàng (siêu âm) — tạo khi doctor chỉ định, số tiền từ bảng giá dịch vụ tại thời điểm chỉ định')
add_bullet('PRESCRIPTION: Tiền thuốc — tạo khi doctor kê đơn + bệnh nhân chọn mua tại phòng khám')

add_para('')
add_para('Trạng thái hoá đơn:', bold=True)
add_bullet('Unpaid → Paid: Staff xác nhận thu tiền mặt (confirmCashPayment), ghi paid_at + paid_by_user_id + payment_method=CASH')
add_bullet('Paid → Refunded: Staff hoàn tiền, ghi refunded_at + refunded_by_user_id + refund_reason')
add_bullet('Unpaid → Cancelled: Khi huỷ lịch/chỉ định, hoá đơn chưa thanh toán bị huỷ theo')

add_para('')
add_para('Quy tắc thanh toán:', bold=True)
add_bullet('Tất cả thanh toán bằng TIỀN MẶT (CASH) tại quầy. KHÔNG có thanh toán điện tử.')
add_bullet('Bệnh nhân KHÔNG được tự đánh dấu đã thanh toán.')
add_bullet('Chống xác nhận trùng: hoá đơn đã Paid thì từ chối, không ghi đè, không cộng doanh thu 2 lần.')
add_bullet('Doanh thu = SUM(invoices.total_amount WHERE status=\'Paid\') - SUM(invoices.total_amount WHERE status=\'Refunded\')')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN VI: VALIDATE LOGIC VÀ RÀNG BUỘC
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN VI: VALIDATE LOGIC VÀ RÀNG BUỘC GIỮA CÁC LUỒNG', 1)

add_heading('6.1. Hàm getStage() — Nguồn chân lý duy nhất', 2)
add_para('Hàm AppointmentStageService.getStage(appointmentId) xác định giai đoạn khám hiện tại, '
         'SUY RA từ dữ liệu gốc mỗi lần gọi (KHÔNG lưu vào DB). Đây là nguồn chân lý DUY NHẤT dùng cho cả UI và server.')
add_para('')
add_para('Logic xét tuần tự từ trên xuống (dừng ở điều kiện đầu tiên khớp):', bold=True)
add_table(
    ['Thứ tự', 'Điều kiện', 'ExamStage', 'Mô tả'],
    [
        ['1', 'appointment.status không phải InProgress và không phải SUCCESS/Completed', 'NOT_STARTED', 'Chưa được bác sĩ tiếp nhận'],
        ['2', 'appointment.status = SUCCESS hoặc Completed', 'FINALIZED', 'Đã hoàn tất'],
        ['3', 'Chưa có medical_record (mr.id IS NULL)', 'CLINICAL_EXAM', 'Cần tạo bệnh án nháp'],
        ['4', 'medical_record.status = Final', 'FINALIZED', 'Bệnh án đã chốt'],
        ['5', 'Có active test_order VÀ có POST_EXAM chưa thanh toán', 'WAITING_PAYMENT', 'Chờ bệnh nhân thanh toán dịch vụ'],
        ['6', 'Có active test_order VÀ tất cả POST_EXAM đã thanh toán', 'WAITING_ULTRASOUND', 'Chờ kết quả siêu âm'],
        ['7', 'Chưa có final_diagnosis', 'ORDER_DECISION', 'Có thể chỉ định CLS hoặc chẩn đoán'],
        ['8', 'Có final_diagnosis VÀ có prescription/treatment_plan', 'READY_TO_FINALIZE', 'Sẵn sàng chốt'],
        ['9', 'Có final_diagnosis nhưng chưa có prescription', 'DIAGNOSIS', 'Sẵn sàng kê đơn'],
    ],
    [1, 7, 3.5, 5.5]
)

add_heading('6.2. Ma trận action được phép theo ExamStage', 2)
add_para('Hàm checkActionAllowed(appointmentId, doctorId, action) chặn ở SERVER:', bold=True)
add_table(
    ['Action', 'ExamStage cho phép', 'Thông báo khi bị chặn (tiếng Việt)'],
    [
        ['startConsultation', 'NOT_STARTED (với Status=Waiting)', '"Không thể tiếp nhận ca: lịch hẹn đang ở giai đoạn «...»"'],
        ['saveClinicalExam', 'CLINICAL_EXAM, ORDER_DECISION', '"Không thể lưu khám lâm sàng: giai đoạn hiện tại là «...»"'],
        ['orderUltrasound', 'ORDER_DECISION', '"Cần lưu bệnh án nháp trước khi chỉ định siêu âm." hoặc "Không thể chỉ định..."'],
        ['cancelUltrasound', 'WAITING_PAYMENT', '"Chỉ có thể huỷ chỉ định khi đang chờ thanh toán. Giai đoạn hiện tại: «...»"'],
        ['saveDiagnosis', 'DIAGNOSIS, READY_TO_FINALIZE', '"Đang chờ kết quả siêu âm..." hoặc "Bệnh nhân chưa thanh toán..."'],
        ['finalizeRecord', 'Không phải WAITING_ULTRASOUND/WAITING_PAYMENT/NOT_STARTED', '"Không thể chốt hồ sơ: còn kết quả siêu âm chưa có." hoặc "chưa thanh toán..."'],
    ],
    [3, 6, 8]
)

add_heading('6.3. Validate đặt lịch (AppointmentValidationService)', 2)
add_para('Tập trung TẤT CẢ validate đặt lịch vào AppointmentValidationService.validateAppointmentInput():')
add_table(
    ['Trường', 'Quy tắc validate'],
    [
        ['Họ tên', 'Không được để trống'],
        ['SĐT', 'Không được để trống; phải bắt đầu bằng 0, 10-11 chữ số'],
        ['Ngày sinh', 'Không được > hôm nay; tuổi 10-100; parse được'],
        ['Bác sĩ', 'Không được để trống'],
        ['Ngày khám', 'Không được để trống; không được trong quá khứ'],
        ['Khung giờ', 'Không được để trống (bắt buộc cho mọi đối tượng)'],
        ['Triệu chứng', 'Bắt buộc 10-500 ký tự; không chỉ chứa số; không có ký tự lặp >5 lần; ít nhất 2 từ; chỉ chứa chữ cái/số/dấu câu cơ bản'],
        ['LMP', 'Không được sau ngày khám; không được > hôm nay; tuổi thai không quá 42 tuần (294 ngày)'],
        ['Trùng lịch', '1 bệnh nhân/1 ngày/1 lịch active (Pending/Confirmed/Waiting/InProgress). Staff có thể override bằng lý do.'],
        ['Slot', 'Phải tồn tại trong doctor_schedules, status=APPROVED, booked_count < max_slots, không trong quá khứ'],
        ['Ca trong ngày', 'Nếu hôm nay: ca chưa kết thúc (end_time > now); nếu <30 phút trước khi kết thúc: từ chối'],
    ],
    [4, 13]
)

add_heading('6.4. Validate check-in (StaffReceptionService.checkInPatient)', 2)
add_para('Điều kiện BẮT BUỘC ĐỦ (kiểm tra tuần tự, lỗi đầu tiên được trả về):')
add_bullet('1. appointment tồn tại')
add_bullet('2. appointment_date = HÔM NAY')
add_bullet('3. status không phải Cancelled')
add_bullet('4. status không phải SUCCESS/Completed')
add_bullet('5. status không phải InProgress (đã check-in rồi?)')
add_bullet('6. status không phải Waiting (đã check-in rồi!)')
add_bullet('7. status không phải Pending (chưa duyệt → "Vui lòng nhấn Duyệt & Tạo Hoá Đơn trước")')
add_bullet('8. Tồn tại PRE_EXAM invoice → nếu không: "Chưa có hóa đơn phí khám"')
add_bullet('9. PRE_EXAM.status = Paid → nếu không: "Bệnh nhân chưa thanh toán phí khám"')
add_bullet('10. Trong khung giờ check-in: không sớm hơn 120 phút trước giờ ca, không sau giờ kết thúc ca')

add_heading('6.5. Validate siêu âm (UltrasoundOrderService)', 2)
add_bullet('isReadyForSonographer: POST_EXAM đã Paid, test_order chưa Completed/Cancelled, chưa có sonographer nhận')
add_bullet('checkSonographerOwnership: chỉ sonographer đã accept ca mới thao tác được')
add_bullet('Tuần tự 4 bước: chưa tải ảnh → không gọi được AI; chưa AI → không ký được; chưa ký → không công bố được')
add_bullet('Nhánh vẽ tay: BẮT BUỘC có ≥1 annotation (validate ở SERVER, không chỉ checkbox client)')
add_bullet('AI timeout: cho phép skip, vẽ thủ công hoàn toàn, ghi log')

add_heading('6.6. Validate thanh toán (StaffReceptionService.confirmCashPayment)', 2)
add_bullet('Chỉ staff role mới được gọi')
add_bullet('appointment.status = Confirmed hoặc Waiting')
add_bullet('PRE_EXAM invoice tồn tại và đang Unpaid (chống xác nhận trùng)')
add_bullet('Cập nhật atomic: status=Paid + paid_at=GETDATE() + paid_by_user_id + payment_method=CASH')
add_bullet('Nếu invoice chưa tồn tại: INSERT mới (UPSERT pattern)')
add_bullet('Backfill invoice_items nếu thiếu')
add_bullet('Gửi NotificationHelper.paymentConfirmed()')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN VII: BẢO MẬT VÀ PHÂN QUYỀN
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN VII: BẢO MẬT VÀ PHÂN QUYỀN', 1)

add_heading('7.1. Xác thực (Authentication)', 2)
add_bullet('AuthenticationFilter: chặn tất cả request không có session hợp lệ (trừ public paths: login, register, forgot-password, google-login, etc.)')
add_bullet('Hỗ trợ 2 phương thức: Username/Password (BCrypt) + Google OAuth 2.0')
add_bullet('Email verification: token gửi qua email, xác nhận trước khi kích hoạt tài khoản')
add_bullet('Password reset: token hết hạn, link qua email')
add_bullet('Session-based: HttpSession lưu user object (id, role, fullName)')

add_heading('7.2. Phân quyền (Authorization)', 2)
add_bullet('AuthorizationFilter: Default Deny — mọi URL không trong whitelist của role sẽ bị chặn (403)')
add_bullet('RBAC: Role → Permissions. 6 roles: Admin, Manager, Doctor, Staff, Sonographer, Patient')
add_bullet('Permission version: bumpPermissionsVersion() khi Admin thay đổi quyền → reload permissions')
add_bullet('Audit log: mọi truy cập (thành công/từ chối) đều ghi lại user, role, URL, IP, timestamp')

add_heading('7.3. Kiểm tra chủ sở hữu (Ownership Check)', 2)
add_para('MỌI action nhận ID từ URL/form đều phải kiểm tra chủ sở hữu/phân công ở SERVER:')
add_bullet('Patient: chỉ xem/sửa lịch hẹn, hoá đơn, bệnh án CỦA CHÍNH MÌNH')
add_bullet('Doctor: chỉ thao tác trên lịch hẹn ĐƯỢC PHÂN CÔNG cho mình')
add_bullet('Sonographer: chỉ thao tác trên ca ĐÃ TIẾP NHẬN')
add_bullet('Staff: chỉ thao tác trong phạm vi lễ tân (không sửa chẩn đoán, không sửa đơn thuốc)')
add_bullet('Sửa ID trên URL để truy cập dữ liệu người khác → BỊ CHẶN + thông báo tiếng Việt')

add_heading('7.4. Bảo vệ dữ liệu', 2)
add_bullet('Email, phone: mã hoá (VARBINARY trong DB, EncryptionUtil)')
add_bullet('Password: BCrypt hash')
add_bullet('CSRF: CsrfFilter kiểm tra token cho các request POST/PUT/DELETE')
add_bullet('Encoding: UTF-8 toàn hệ thống (EncodingFilter)')
add_bullet('Chống gửi trùng (idempotency): bấm 2 lần/F5 không tạo bản ghi trùng, không cộng/trừ số liệu 2 lần')
add_bullet('Transaction: mọi thao tác ghi nhiều bảng đều trong transaction có rollback')
add_bullet('Audit log: mọi thao tác đổi trạng thái + đụng tiền đều ghi audit log')

add_heading('7.5. Bảo vệ dữ liệu bệnh nhân', 2)
add_bullet('Bệnh nhân CHỈ xem được hồ sơ bệnh án đã chốt (Final). Bản nháp (Draft) tuyệt đối không lộ.')
add_bullet('Kiểm tra chủ sở hữu ở TẤT CẢ màn hình: appointments, invoices, medical_records, ultrasound_reports')
add_bullet('Ảnh siêu âm lưu bytes gốc, KHÔNG re-encode qua ImageIO (tránh giảm chất lượng ảnh chẩn đoán)')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN VIII: QUY TẮC NGHIỆP VỤ ĐÃ CHỐT
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN VIII: QUY TẮC NGHIỆP VỤ ĐÃ CHỐT', 1)
add_table(
    ['#', 'Vấn đề', 'Quyết định đã chốt'],
    [
        ['1', 'Đổi lịch (Change Appointment)', 'Chỉ đổi slot CÙNG MỘT BÁC SĨ. Muốn đổi bác sĩ → huỷ rồi đặt lại.'],
        ['2', 'Đổi/huỷ khi đã thanh toán', 'Bệnh nhân KHÔNG tự làm được. Chuyển sang quầy, staff xử lý hoàn tiền.'],
        ['3', 'Hạn chót đổi/huỷ', '2 giờ trước giờ bắt đầu ca.'],
        ['4', 'Kê đơn có chặn chốt hồ sơ không', 'KHÔNG chặn. Hoá đơn thuốc do staff xử lý riêng ở quầy.'],
        ['5', 'Nhiều chỉ định siêu âm', 'CHO PHÉP. Mở khoá chẩn đoán khi TẤT CẢ đã hoàn thành.'],
        ['6', 'AI Engine chết', 'Cho phép BỎ QUA AI, vẽ thủ công hoàn toàn, ghi log ca không qua AI.'],
        ['7', 'Thống kê vs Doanh thu', 'Thống kê = số lượt. Doanh thu = tiền. HAI MÀN HÌNH TÁCH BIỆT.'],
        ['8', 'Đổi giá dịch vụ', 'Hoá đơn đã phát hành KHÔNG ĐỔI. Giá mới chỉ áp cho chỉ định sau đó.'],
        ['9', 'Ca kéo dài qua ngày', 'CHẤP NHẬN, không auto huỷ. Manager có danh sách ca chưa chốt quá 24h.'],
        ['10', 'Bác sĩ khám song song', 'CHO PHÉP nhiều ca InProgress cùng lúc.'],
        ['11', 'Thanh toán', 'TIỀN MẶT tại quầy. KHÔNG cổng thanh toán điện tử.'],
        ['12', 'Bệnh nhân tự đánh dấu thanh toán', 'TUYỆT ĐỐI KHÔNG.'],
        ['13', 'BS lâm sàng sửa kết luận siêu âm', 'KHÔNG ĐƯỢC. Kết luận siêu âm là của sonographer.'],
        ['14', 'Staff sửa chẩn đoán/đơn thuốc', 'KHÔNG ĐƯỢC. Chẩn đoán và đơn thuốc là của bác sĩ.'],
        ['15', 'Trạng thái treo vĩnh viễn', 'KHÔNG có trạng thái nào gây treo vĩnh viễn.'],
        ['16', 'Giá khám', 'Lưu cứng vào base_fee tại thời điểm đặt lịch. Không tính lại.'],
        ['17', 'Giới hạn đổi lịch', 'Tối đa 2 lần/lịch hẹn. Đếm từ audit log, không thêm cột.'],
        ['18', 'Đặt lịch cùng ngày', '1 bệnh nhân/1 ngày/1 lịch active. Staff có thể override với lý do.'],
        ['19', 'Bác sĩ hết ca', 'Tự động chuyển bệnh nhân đang chờ sang ca tiếp theo của bác sĩ.'],
    ],
    [0.5, 5, 11.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PHẦN IX: DANH SÁCH SERVLET & API
# ══════════════════════════════════════════════════════════════════════
add_heading('PHẦN IX: DANH SÁCH SERVLET / API REFERENCE', 1)

add_heading('9.1. Patient Servlets', 2)
add_table(
    ['Servlet', 'URL Pattern', 'Method', 'Mô tả'],
    [
        ['PatientBookingServlet', '/patient/booking', 'GET/POST', 'Đặt lịch online'],
        ['PatientSlotApiServlet', '/patient/slots', 'GET', 'API JSON: danh sách slot khả dụng'],
        ['PatientAppointmentServlet', '/patient/appointments', 'GET/POST', 'Danh sách lịch hẹn + huỷ/đổi lịch'],
        ['PatientMedicalRecordServlet', '/patient/medical-record', 'GET', 'Xem hồ sơ bệnh án (đã chốt)'],
        ['PatientInvoicesServlet', '/patient/invoices', 'GET', 'Danh sách hoá đơn'],
        ['PatientPaymentServlet', '/patient/payment', 'GET', 'Chi tiết thanh toán'],
        ['PatientProfileServlet', '/patient/profile', 'GET/POST', 'Quản lý hồ sơ cá nhân'],
        ['PatientReviewServlet', '/patient/review', 'POST', 'Đánh giá bác sĩ'],
    ],
    [4, 4, 1.8, 7]
)

add_heading('9.2. Staff Servlets', 2)
add_table(
    ['Servlet', 'URL Pattern', 'Method', 'Mô tả'],
    [
        ['StaffQueueServlet', '/staff/queue', 'GET/POST', 'Hàng đợi + approve + check-in + cancel + mark-priority + no-show + confirm-payment + refund'],
        ['StaffBookingServlet', '/staff/booking', 'GET/POST', 'Tạo lịch thủ công'],
        ['StaffEditServlet', '/staff/edit', 'GET/POST', 'Sửa thông tin lịch hẹn'],
        ['StaffDoctorScheduleServlet', '/staff/schedules', 'GET', 'Xem lịch làm việc bác sĩ'],
        ['StaffProfileServlet', '/staff/profile', 'GET/POST', 'Hồ sơ cá nhân'],
        ['StaffPatientLookupServlet', '/staff/patient-lookup', 'GET', 'Tìm bệnh nhân theo SĐT'],
    ],
    [4, 4, 1.8, 7]
)

add_heading('9.3. Doctor Servlets', 2)
add_table(
    ['Servlet', 'URL Pattern', 'Method', 'Mô tả'],
    [
        ['DoctorDashboardServlet', '/doctor/dashboard', 'GET', 'Dashboard'],
        ['DoctorAppointmentServlet', '/doctor/appointments', 'GET/POST', 'DS bệnh nhân + tiếp nhận ca'],
        ['DoctorPatientListServlet', '/doctor/patients', 'GET', 'DS bệnh nhân'],
        ['DoctorPatientHistoryServlet', '/doctor/patient-history', 'GET', 'Lịch sử bệnh nhân'],
        ['MedicalRecordServlet', '/doctor/medical-record', 'GET/POST', 'Hồ sơ bệnh án (5 khối)'],
        ['DoctorUltrasoundRequestServlet', '/doctor/ultrasound-request', 'POST', 'Chỉ định/huỷ siêu âm'],
        ['DoctorResultsServlet', '/doctor/results', 'GET', 'Xem kết quả siêu âm'],
        ['DoctorPrescriptionListServlet', '/doctor/prescriptions', 'GET', 'DS đơn thuốc'],
        ['PrescriptionServlet', '/doctor/prescription', 'POST', 'Tạo đơn thuốc'],
        ['DoctorScheduleServlet', '/doctor/schedule', 'GET/POST', 'Đăng ký lịch làm việc'],
        ['DoctorProfileServlet', '/doctor/profile', 'GET/POST', 'Hồ sơ cá nhân'],
        ['DoctorPregnancyServlet', '/doctor/pregnancy', 'GET/POST', 'Quản lý thai kỳ'],
    ],
    [4, 4, 1.8, 7]
)

add_heading('9.4. Sonographer Servlets', 2)
add_table(
    ['Servlet', 'URL Pattern', 'Method', 'Mô tả'],
    [
        ['UltrasoundWaitingListServlet', '/sonographer/waiting-list', 'GET', 'Dashboard DS ca chờ'],
        ['UltrasoundDetailServlet', '/sonographer/detail', 'GET/POST', 'Chi tiết ca + accept'],
        ['UltrasoundUploadServlet', '/sonographer/upload', 'POST', 'Upload ảnh siêu âm'],
        ['UltrasoundAnalyzeServlet', '/sonographer/analyze', 'POST', 'Gửi phân tích AI'],
        ['UltrasoundModelServlet', '/sonographer/model', 'GET/POST', 'Xử lý kết quả AI (Accept/Reject)'],
        ['MockAiEngineServlet', '/sonographer/mock-ai', 'POST', 'AI Engine mock (loopback only)'],
        ['AiImageStreamServlet', '/sonographer/ai-image', 'GET', 'Stream ảnh kết quả AI'],
        ['UltrasoundImageStreamServlet', '/sonographer/image', 'GET', 'Stream ảnh siêu âm gốc'],
        ['SonographerProfileServlet', '/sonographer/profile', 'GET/POST', 'Hồ sơ cá nhân'],
    ],
    [4, 4, 1.8, 7]
)

add_heading('9.5. Manager Servlets', 2)
add_table(
    ['Servlet', 'URL Pattern', 'Method', 'Mô tả'],
    [
        ['DashboardServlet', '/manager/dashboard', 'GET', 'Dashboard'],
        ['ManagerServiceServlet', '/manager/services', 'GET/POST', 'Quản lý dịch vụ + giá'],
        ['ManagerRevenueServlet', '/manager/revenue', 'GET', 'Báo cáo doanh thu'],
        ['ManagerStatisticsServlet', '/manager/statistics', 'GET', 'Thống kê số lượt'],
        ['ManagerScheduleServlet', '/manager/schedules', 'GET/POST', 'Duyệt lịch làm việc'],
        ['ManagerTimeSlotServlet', '/manager/slots', 'GET/POST', 'Quản lý khung giờ'],
        ['ManagerDoctorServlet', '/manager/doctors', 'GET', 'DS bác sĩ'],
        ['ManagerMedicineServlet', '/manager/medicines', 'GET/POST', 'Quản lý thuốc'],
        ['ManagerProfileServlet', '/manager/profile', 'GET/POST', 'Hồ sơ cá nhân'],
    ],
    [4, 4, 1.8, 7]
)

add_heading('9.6. Admin Servlets', 2)
add_table(
    ['Servlet', 'URL Pattern', 'Method', 'Mô tả'],
    [
        ['AdminUserServlet', '/admin/users', 'GET/POST', 'Quản lý người dùng'],
        ['AdminRoleServlet', '/admin/roles', 'GET/POST', 'Quản lý vai trò & phân quyền'],
        ['AdminMedicineServlet', '/admin/medicines', 'GET/POST', 'Quản lý danh mục thuốc'],
        ['AdminServiceServlet', '/admin/services', 'GET/POST', 'Quản lý dịch vụ'],
        ['AdminPriceServlet', '/admin/pricing', 'GET/POST', 'Quản lý bảng giá'],
        ['AdminAuditLogServlet', '/admin/audit-logs', 'GET', 'Xem audit log'],
        ['AdminProfileServlet', '/admin/profile', 'GET/POST', 'Hồ sơ cá nhân'],
    ],
    [4, 4, 1.8, 7]
)

add_heading('9.7. Auth Servlets (Public)', 2)
add_table(
    ['Servlet', 'URL Pattern', 'Method', 'Mô tả'],
    [
        ['LoginServlet', '/login', 'GET/POST', 'Đăng nhập username/password'],
        ['GoogleLoginServlet', '/google-login', 'GET', 'Google OAuth redirect'],
        ['GoogleServerLoginServlet', '/google-server-login', 'GET', 'Google OAuth callback'],
        ['RegisterServlet', '/register', 'GET/POST', 'Đăng ký tài khoản'],
        ['VerifyEmailServlet', '/verify-email', 'GET', 'Xác nhận email'],
        ['ForgotPasswordServlet', '/forgot-password', 'GET/POST', 'Quên mật khẩu'],
        ['ResetPasswordServlet', '/reset-password', 'GET/POST', 'Đặt lại mật khẩu'],
        ['ChangePasswordServlet', '/change-password', 'POST', 'Đổi mật khẩu'],
        ['LogoutServlet', '/logout', 'GET', 'Đăng xuất'],
    ],
    [4, 4, 1.8, 7]
)

# ══════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════
output_path = r'c:\Users\admin\IdeaProjects\Clinic-Appointment-Management-System\docs\CAMS_RDS_TaiLieuDacTa.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print('Done!')
