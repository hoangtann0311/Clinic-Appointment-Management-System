#!/usr/bin/env python3
"""
Generate RDS Document for CAMS following Template2_RDS structure.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 51, 102)
    return h

def P(text, bold=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    r.bold = bold
    return p

def B(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def T(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]
            c.text = str(val) if val else ""
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Requirement & Design Specification')
r.font.size = Pt(24); r.font.bold = True; r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Clinic Appointment Management System (CAMS)')
r.font.size = Pt(18); r.font.name = 'Times New Roman'

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run('Version: 1.0')
r.font.size = Pt(14); r.font.name = 'Times New Roman'

for _ in range(4): doc.add_paragraph()
loc = doc.add_paragraph()
loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = loc.add_run(f'– Hanoi, {datetime.date.today().strftime("%B %Y")} –')
r.font.size = Pt(13); r.font.name = 'Times New Roman'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# RECORD OF CHANGES
# ═══════════════════════════════════════════════════════════════
H('Record of Changes', 1)
P('*A - Added  M - Modified  D - Deleted')
T(['Date', 'Version', 'Description', 'A/M/D', 'Author'],
  [[datetime.date.today().strftime('%d/%m/%Y'), '1.0', 'Initial RDS document for CAMS', 'A', 'CAMS Team']],
  [3,2,7,1.5,3])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
P('Contents')
toc = [
    ('I. Overview', [
        '1. Context Diagram',
        '2. Main Business Processes',
        '  2.1 Create Appointment Flow',
        '  2.2 Clinical Examination & Ultrasound Order Flow',
        '  2.3 Ultrasound Flow',
        '  2.4 Diagnosis and Prescription Flow',
        '  2.5 Payment & Check-in Flow',
        '3. User Requirements',
        '  3.1 Actors',
        '  3.2 Use Cases',
        '4. Overall Functionalities',
        '  4.1 Screens Flow',
        '  4.2 Screen Descriptions',
        '  4.3 Screen Authorization',
        '  4.4 Non-UI Functions',
        '5. System High Level Design',
        '  5.1 Database Design',
        '  5.2 Code Packages',
    ]),
    ('II. Requirement Specifications', [
        '1. Authentication (UC-01 to UC-05)',
        '2. Appointment Booking (UC-06 to UC-08)',
        '3. Reception & Queue Management (UC-09 to UC-17)',
        '4. Clinical Examination (UC-18 to UC-22)',
        '5. Ultrasound Diagnostics (UC-23 to UC-30)',
        '6. Diagnosis & Prescription (UC-31 to UC-34)',
        '7. Payment Management (UC-35 to UC-38)',
        '8. Medical Records (UC-39 to UC-41)',
        '9. Patient Portal (UC-42 to UC-46)',
        '10. Service & Price Management (UC-47 to UC-50)',
        '11. Doctor Schedule Management (UC-51 to UC-55)',
        '12. Revenue & Statistics (UC-56 to UC-59)',
        '13. User & Role Management (UC-60 to UC-64)',
        '14. Pharmacy Management (UC-65 to UC-68)',
        '15. System & Security (UC-69 to UC-72)',
    ]),
    ('III. Design Specifications', [
        '1. Authentication',
        '2. Appointment Booking',
        '3. Reception & Queue Management',
        '4. Clinical Examination & Medical Record',
        '5. Ultrasound Diagnostics & AI Integration',
        '6. Diagnosis & Prescription',
        '7. Payment Management',
        '8. Revenue & Statistics',
        '9. Service & Price Management',
        '10. User, Role & Security Management',
    ]),
]
for section, items in toc:
    P(section, bold=True, size=13)
    for item in items:
        P(item, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# I. OVERVIEW
# ═══════════════════════════════════════════════════════════════
H('I. Overview', 1)

# 1. Context Diagram
H('1. Context Diagram', 2)
P('Diagram 1: Context diagram of CAMS — Clinic Appointment Management System')
P('[Context Diagram — showing 6 actors (Admin, Manager, Patient, Doctor, Sonographer, Staff) interacting with CAMS central system via HTTP/Web browser. External services: Google OAuth for social login, Python AI Engine for ultrasound image analysis, Email SMTP for notifications.]')
doc.add_paragraph()

# 2. Main Business Processes
H('2. Main Business Processes', 2)

H('2.1 Create Appointment Flow', 3)
P('Diagram 2: Swimlane — Create Appointment Flow')
P('Actors: Patient, Staff, System')
P('Description: Patient books appointment online (selects doctor → views price → picks available slot → enters symptoms/LMP → confirms). Staff creates manual booking (looks up patient by phone → selects doctor → picks slot → enters symptoms → confirms). System validates: slot availability, no duplicate same-day booking, symptoms 10-500 chars, LMP validity, age 10-100, phone format. Result: Appointment (Pending) with base_fee locked, booked_count incremented.')
doc.add_paragraph()

H('2.2 Clinical Examination & Ultrasound Order Flow', 3)
P('Diagram 3: Swimlane — Clinical Examination & Ultrasound Order Flow')
P('Actors: Doctor, Staff, System')
P('Description: Doctor accepts patient (Waiting→InProgress), fills clinical exam form (vitals, symptoms, history), saves draft. Doctor decides: no imaging → proceed to diagnosis; or order ultrasound → selects service → enters mandatory reason (10-500 chars) → creates test_order + POST_EXAM invoice. System locks diagnosis/finalize blocks, shows waiting reason in Vietnamese. Staff collects POST_EXAM payment (cash). System unlocks ultrasound for sonographer.')
doc.add_paragraph()

H('2.3 Ultrasound Flow', 3)
P('Diagram 4: Swimlane — Ultrasound Flow')
P('Actors: Sonographer, AI Engine, System')
P('Description: Sonographer accepts case → uploads images (raw bytes, no re-encode). Sends to AI Engine (HTTP with timeout). Two branches: Accept AI Result or Reject AI (mandatory manual annotation). Fill findings → Sign → Publish. System validates: sequential 4-step flow, at least 1 annotation if manual draw, ownership check. On publish: test_order→Completed, report saved, DOCTOR NOTIFIED (critical handoff). AI failure: skip AI, manual draw only, log entry.')
doc.add_paragraph()

H('2.4 Diagnosis and Prescription Flow', 3)
P('Diagram 5: Swimlane — Diagnosis and Prescription Flow')
P('Actors: Doctor, System')
P('Description: Doctor views ultrasound results (read-only: images, AI result, sonographer conclusion). Confirms viewing. Enters diagnosis and conclusion. Prescribes medication if needed, patient chooses in-clinic purchase or external. System validates: no pending ultrasound orders blocking diagnosis. Finalize: medical_record→Final, appointment→SUCCESS, queue_number cleared, patient notified. Prescription invoice handled separately by staff — does NOT block finalization.')
doc.add_paragraph()

H('2.5 Payment & Check-in Flow', 3)
P('Diagram 6: Swimlane — Payment & Check-in Flow')
P('Actors: Staff, System')
P('Description: Staff approves booking (Pending→Confirmed) → creates PRE_EXAM invoice (Unpaid, copies base_fee). Patient pays cash at counter → Staff confirms payment (Paid, records paid_at/paid_by_user_id/payment_method=CASH). Staff checks in (Confirmed→Waiting, assigns queue_number) — requires: today, PRE_EXAM Paid, within time window. Staff can mark priority (reason 5-500 chars, renumbers queue). Staff can mark NoShow (after shift end) or Cancel (with refund if paid).')
doc.add_paragraph()

doc.add_page_break()

# 3. User Requirements
H('3. User Requirements', 2)

H('3.1 Actors', 3)
T(['Actor', 'Role', 'Description'],
  [
      ['Admin', 'System Administrator', 'Manages users, roles, permissions. Manages medicines and services catalog. Views audit logs and system security settings.'],
      ['Manager', 'Clinic Manager', 'Manages medical services & pricing. Views revenue reports and service statistics. Approves/rejects doctor work schedules. Manages time slots. Monitors stuck appointments >24h.'],
      ['Patient', 'Patient (Sản phụ)', 'Books appointments online. Views own appointments with status and queue number. Views finalized medical records. Cancels/reschedules appointments (with constraints). Views payment history. Reviews doctors.'],
      ['Doctor', 'Clinical Doctor (Bác sĩ lâm sàng)', 'Views today\'s patient list. Accepts patients for examination. Fills clinical exam forms. Orders/cancels ultrasound. Views ultrasound results (read-only). Diagnoses and prescribes medications. Finalizes medical records. Manages personal work schedules.'],
      ['Sonographer', 'Ultrasound Doctor (Bác sĩ siêu âm)', 'Views ultrasound waiting list. Accepts ultrasound cases. Uploads ultrasound images. Sends for AI analysis. Accepts/rejects AI results. Draws manual annotations. Signs and publishes results.'],
      ['Staff', 'Receptionist (Nhân viên lễ tân)', 'Views reception queue. Approves bookings & creates invoices. Confirms cash payments. Checks in patients. Marks priority/no-show. Cancels appointments with refund handling. Creates manual bookings for walk-in patients. Edits pending appointments.'],
  ],
  [2.5, 3, 11.5])

H('3.2 Use Cases', 2)
H('a. Diagram(s)', 4)
P('Diagram 7: Use Case Diagram — CAMS System')
P('[Use Case Diagram showing 6 actors with their respective use cases. Patient: Book Appointment, View Appointments, Cancel Appointment, Change Appointment, View Medical Record, View Payment History, Review Doctor. Staff: View Reception Queue, Approve Booking, Confirm Payment, Check In Patient, Mark Priority, Cancel Appointment, Mark NoShow, Create Manual Booking, Update Appointment, Refund Invoice. Doctor: View Today\'s Appointments, Update Medical Record, Order Ultrasound, Cancel Ultrasound Order, Diagnose and Prescribe, Finalize Medical Record, Manage Work Schedules. Sonographer: Accept Case, Upload Images, Analyze with AI, Accept/Reject AI, Draw Manually, Sign and Confirm, Publish Results. Manager: Manage Services, View Price History, Manage Revenue, View Statistics, Approve Schedules, Manage Slots, Manage Doctors. Admin: Manage Users, Manage Roles, Manage Medicines, Manage Services, View Audit Logs.]')

doc.add_paragraph()

H('b. Descriptions', 4)
P('Detailed use case descriptions are provided in Part II — Requirement Specifications below.', bold=True)

doc.add_page_break()

# 4. Overall Functionalities
H('4. Overall Functionalities', 2)

H('4.1 Screens Flow', 3)
P('[Screens Flow Diagram — showing navigation paths between all screens in the system, grouped by actor:]')
screen_flows = [
    ('Authentication', 'Login → Register → Forgot Password → Reset Password → Verify Email → Home Dashboard'),
    ('Patient', 'Home → Book Appointment → Appointment List → Medical Record Detail → Invoice List → Payment Detail → Profile → Review Doctor'),
    ('Staff', 'Reception Queue → Manual Booking → Edit Appointment → Doctor Schedules → Patient Lookup → Profile'),
    ('Doctor', 'Dashboard → Patient List → Medical Record Form (5 blocks) → Ultrasound Request → Results View → Prescription List → Schedule Management → Profile'),
    ('Sonographer', 'Waiting List → Case Detail → Upload Images → AI Analysis → AI Model (Accept/Reject) → Sign → Publish → Profile'),
    ('Manager', 'Dashboard → Services → Service Detail → Price History → Revenue → Revenue Detail → Statistics → Schedules → Slots → Doctors → Medicines → Profile'),
    ('Admin', 'Dashboard → Users → Roles → Medicines → Services → Pricing → Audit Logs → Profile'),
]
for title, desc in screen_flows:
    B(f'{title}: {desc}')

doc.add_paragraph()

H('4.2 Screen Descriptions', 3)
P('Key screens in the CAMS system:')
T(['Screen', 'Actor', 'Description', 'Key Validations'],
  [
      ['Login / Register', 'All (Public)', 'Login with username/password or Google OAuth. Register new patient account with email verification.', 'Email format, password strength, duplicate email/phone check, Google token validation'],
      ['Book Appointment', 'Patient', 'Step-by-step booking: select doctor→view price→pick date→choose slot→enter symptoms/LMP→confirm. Shows total price with "pay at counter" notice.', 'Slot availability, no duplicate same-day, symptoms 10-500 chars, LMP validity, age 10-100, phone 0xxxxxxxxx'],
      ['Reception Queue', 'Staff', 'Paginated list of today\'s appointments. Filter by status/search. Actions: approve, confirm payment, check-in, mark priority, cancel, no-show, refund.', 'Status-based action eligibility, PRE_EXAM must be paid before check-in, today-only for check-in/no-show, shift time window'],
      ['Medical Record Form', 'Doctor', '5 accordion blocks: Accept→Clinical Exam→Imaging Decision→Diagnosis→Finalize. Blocks unlock sequentially per getStage(). Read-only ultrasound results.', 'getStage() blocks wrong actions, ownership check, mandatory reason 10-500 for ultrasound order'],
      ['Ultrasound Detail', 'Sonographer', '4-step flow: Accept case→Upload images→AI Analysis (or manual draw)→Sign & Publish.', 'Sequential step enforcement, ≥1 manual annotation if rejecting AI, ownership check, AI timeout handling'],
      ['Revenue Report', 'Manager', 'Revenue by date range, broken down by PRE_EXAM/POST_EXAM/PRESCRIPTION. Read directly from invoices (Paid - Refunded).', 'No aggregate table, real-time query'],
      ['Service Management', 'Manager', 'CRUD services + price. Changing price creates price_history entry. Old invoices unaffected.', 'Cannot delete used services (deactivate only), new price only for future orders'],
  ],
  [3, 1.5, 7.5, 5])

doc.add_page_break()

H('4.3 Screen Authorization', 3)
P('Role-based access control matrix (down to screen level):')
T(
    ['Screen / Feature', 'Admin', 'Manager', 'Doctor', 'Sonographer', 'Staff', 'Patient'],
    [
        ['Login / Register', '✓', '✓', '✓', '✓', '✓', '✓'],
        ['Admin Dashboard', '✓', '-', '-', '-', '-', '-'],
        ['User Management', '✓', '-', '-', '-', '-', '-'],
        ['Role & Permission Mgmt', '✓', '-', '-', '-', '-', '-'],
        ['Audit Logs', '✓', '-', '-', '-', '-', '-'],
        ['Manager Dashboard', '-', '✓', '-', '-', '-', '-'],
        ['Service & Price Mgmt', '✓', '✓', '-', '-', '-', '-'],
        ['Price History', '-', '✓', '-', '-', '-', '-'],
        ['Revenue Report', '-', '✓', '-', '-', '-', '-'],
        ['Service Statistics', '-', '✓', '-', '-', '-', '-'],
        ['Schedule Approval', '-', '✓', '-', '-', '-', '-'],
        ['Time Slot Mgmt', '-', '✓', '-', '-', '-', '-'],
        ['Doctor Management', '-', '✓', '-', '-', '-', '-'],
        ['Medicine Catalog', '✓', '✓', '-', '-', '-', '-'],
        ['Reception Queue', '-', '-', '-', '-', '✓', '-'],
        ['Manual Booking', '-', '-', '-', '-', '✓', '-'],
        ['Confirm Payment', '-', '-', '-', '-', '✓', '-'],
        ['Check-in / Priority', '-', '-', '-', '-', '✓', '-'],
        ['Cancel / NoShow / Refund', '-', '-', '-', '-', '✓', '-'],
        ['Doctor Dashboard', '-', '-', '✓', '-', '-', '-'],
        ['Patient List (Today)', '-', '-', '✓', '-', '-', '-'],
        ['Medical Record (Edit)', '-', '-', '✓', '-', '-', '-'],
        ['Medical Record (Read)', '-', '-', '✓', '✓', '-', '✓ (final only)'],
        ['Ultrasound Order', '-', '-', '✓', '-', '-', '-'],
        ['Prescription Mgmt', '-', '-', '✓', '-', '-', '-'],
        ['Schedule Registration', '-', '-', '✓', '-', '-', '-'],
        ['Sonographer Dashboard', '-', '-', '-', '✓', '-', '-'],
        ['Ultrasound Case Detail', '-', '-', '-', '✓', '-', '-'],
        ['Upload/AI/Annotation', '-', '-', '-', '✓', '-', '-'],
        ['Sign & Publish', '-', '-', '-', '✓', '-', '-'],
        ['Book Appointment', '-', '-', '-', '-', '-', '✓'],
        ['My Appointments', '-', '-', '-', '-', '-', '✓'],
        ['My Medical Records', '-', '-', '-', '-', '-', '✓ (final only)'],
        ['My Invoices', '-', '-', '-', '-', '-', '✓'],
        ['Cancel/Reschedule', '-', '-', '-', '-', '-', '✓ (with constraints)'],
        ['Profile (own)', '-', '✓', '✓', '✓', '✓', '✓'],
    ],
    [5.5, 1.3, 1.3, 1.3, 1.8, 1.3, 1.3]
)

doc.add_page_break()

H('4.4 Non-UI Functions', 3)
P('System functions without direct UI (batch jobs, services, APIs, filters):')
T(['Function', 'Type', 'Description'],
  [
      ['AuthenticationFilter', 'Servlet Filter', 'Intercepts all requests. Checks valid session. Redirects to login for unauthenticated access. Passes through public paths (login, register, forgot-password, google-login).'],
      ['AuthorizationFilter', 'Servlet Filter', 'Default-deny RBAC filter. Whitelists URLs per role. Logs all access (success/denied) to audit log. Supports dynamic permission version bumping. Returns 403 on denial.'],
      ['CsrfFilter', 'Servlet Filter', 'Anti-CSRF token validation for POST/PUT/DELETE requests.'],
      ['EncodingFilter', 'Servlet Filter', 'UTF-8 encoding for all requests/responses.'],
      ['SlotHoldExpiryListener', 'Servlet Listener', 'Cleans up expired slot holds (prevent race condition during booking).'],
      ['MockAiEngineServlet', 'HTTP API (internal)', 'Mock AI engine endpoint for ultrasound image analysis. IP-restricted (loopback only). Returns detection results with bounding boxes and confidence scores. Must be replaced with token-based auth behind reverse proxy.'],
      ['AiPredictionService', 'Service', 'HTTP client calling external Python AI Engine for ultrasound analysis. Configurable timeout. Handles connection errors gracefully (allows skip-AI fallback).'],
      ['NotificationHelper', 'Utility Service', 'Centralized notification creation. Sends notifications for: new appointment, schedule approved/rejected, payment confirmed, ultrasound results ready, record finalized, refund processed. All notifications stored in DB for in-app display.'],
      ['AuditUtil / AuditLogDAO', 'Utility Service', 'Audit logging for ALL status changes and financial operations. Records: user_id, action, table_name, old_value, new_value, ip_address, timestamp.'],
      ['AppointmentStageService.getStage()', 'Service', 'Pure function determining exam stage from raw data. No DB writes. Used by both UI (show/hide blocks) and server (block invalid actions). Single source of truth for exam workflow state.'],
      ['StaffReceptionService.autoMovePatientsToNextShift()', 'Auto-job (on page load)', 'Triggered on reception page load. Detects ended shifts (end_time ≤ now). Moves waiting patients to doctor\'s next available shift. Transactional with audit log and patient notification.'],
      ['GoogleAuthService', 'Service', 'Google OAuth 2.0 integration. Handles token exchange, user info retrieval. Auto-provisions new patient accounts from Google profile.'],
      ['PasswordService', 'Service', 'Password hashing (BCrypt), verification, strength validation. Password reset token generation and validation.'],
      ['EncryptionUtil', 'Utility', 'Encrypts sensitive PII fields (email, phone) stored as VARBINARY in database.'],
  ],
  [4, 2, 11])

doc.add_page_break()

# 5. System High Level Design
H('5. System High Level Design', 2)

H('5.1 Database Design', 3)
H('a. Database Schema', 4)
P('[Database Schema Diagram — 32 tables in ObstetricsClinicDB database. Key relationships shown below.]')
P('Core entity relationships:')
P('users (1)──(1) patients/doctors/sonographers', bold=True)
P('patients (1)──(N) appointments (N)──(1) doctors', bold=True)
P('appointments (1)──(1) medical_records (1)──(N) test_orders (N)──(1) services', bold=True)
P('medical_records (1)──(N) prescriptions (1)──(N) prescription_items (N)──(1) medicines', bold=True)
P('test_orders (1)──(N) ultrasound_images (1)──(N) ai_analysis_results', bold=True)
P('test_orders (1)──(N) ultrasound_annotations, (1)──(N) ultrasound_reports', bold=True)
P('appointments (1)──(N) invoices (1)──(N) invoice_items', bold=True)
P('doctors (1)──(N) doctor_schedules (N)──(1) shifts', bold=True)
P('roles (1)──(N) role_permissions (N)──(1) permissions', bold=True)

doc.add_paragraph()
H('b. Table Descriptions', 4)
T(['#', 'Table', 'Key Columns', 'Description'],
  [
      ['1', 'users', 'id, username, email(encrypted), password_hash, phone(encrypted), role_id, status, is_verified, google_id, auth_provider, is_deleted, created_at', 'User accounts with auth info. Supports local + Google OAuth.'],
      ['2', 'roles', 'id, role_name, description', '6 roles: Admin, Manager, Doctor, Staff, Sonographer, Patient'],
      ['3', 'permissions', 'id, permission_key, permission_name, module, description', 'Fine-grained permissions for RBAC'],
      ['4', 'role_permissions', 'role_id, permission_id', 'M:N mapping roles to permissions'],
      ['5', 'patients', 'id, user_id, full_name, phone_number, date_of_birth, address, cccd', 'Patient profiles linked to user account'],
      ['6', 'doctors', 'id, user_id, full_name, specialization, degree, experience_years, avatar_url, bio', 'Doctor profiles. experience_years determines base exam fee'],
      ['7', 'sonographers', 'id, user_id, experience_years, qualification, room_no, status', 'Sonographer profiles'],
      ['8', 'appointments', 'id, patient_id, doctor_id, appointment_date, status, symptoms, LMP, base_fee, queue_number, is_priority, schedule_id, time_slot, held_by, held_until', 'CORE TABLE. Status: Pending→Confirmed→Waiting→InProgress→SUCCESS. base_fee locked at booking time.'],
      ['9', 'medical_records', 'id, appointment_id, clinical_notes, final_diagnosis, status, weight_kg, blood_pressure, pulse_bpm, temperature_c, gestational_age_weeks, fetal_heart_rate, risk_flags_json, treatment_plan', 'Clinical examination data. Status: Draft→Final. 31 columns including obstetric-specific fields.'],
      ['10', 'test_orders', 'id, medical_record_id, doctor_id, service_id, status, reorder_reason, sonographer_user_id, accepted_at', 'Ultrasound orders. Status: Pending→Ordered→InProgress→Completed/Cancelled.'],
      ['11', 'prescriptions', 'id, medical_record_id, prescription_code, status, purchase_decision, purchase_decided_at', 'Prescriptions. purchase_decision: AT_CLINIC or EXTERNAL.'],
      ['12', 'prescription_items', 'id, prescription_id, medicine_id, quantity, dosage', 'Line items in a prescription'],
      ['13', 'medicines', 'id, name, price, unit, medicine_code, dosage, stock_quantity, is_active, category_id', 'Medicine catalog managed by Admin/Manager'],
      ['14', 'medicine_categories', 'id, category_name, description, is_active', 'Medicine category groupings'],
      ['15', 'medicine_price_history', 'id, medicine_id, old_price, new_price, change_reason, changed_by, created_at', 'Audit trail for medicine price changes'],
      ['16', 'invoices', 'id, appointment_id, total_amount, status, transaction_code, invoice_type, payment_method, confirmed_by, confirmed_at, paid_at, paid_by_user_id, payment_note', 'Invoices. Types: PRE_EXAM, POST_EXAM, PRESCRIPTION. Status: Unpaid→Paid→Refunded. paid_at/paid_by_user_id/payment_method provide cash payment evidence.'],
      ['17', 'invoice_items', 'id, invoice_id, item_type, item_id, quantity, unit_price, subtotal', 'Invoice line items'],
      ['18', 'services', 'id, service_name, price, service_code, category_id, duration_mins, requires_fasting, is_active', 'Medical services catalog (ultrasound types, etc.)'],
      ['19', 'service_categories', 'id, category_name, description, is_active', 'Service category groupings'],
      ['20', 'price_history', 'id, service_id, old_price, new_price, change_reason, changed_by, created_at', 'Audit trail for service price changes'],
      ['21', 'doctor_schedules', 'id, doctor_id, work_date, shift_id, max_slots, booked_count, is_approved, status, rejection_reason, approved_by, created_by, notes, version', 'Doctor work schedules. Replaces old time_slots. booked_count updated atomically. version for optimistic locking.'],
      ['22', 'shifts', 'id, name, start_time, end_time, description, is_active', 'Work shifts (e.g., Ca sáng 7:00-12:00, Ca chiều 13:00-17:00)'],
      ['23', 'ultrasound_images', 'id, test_order_id, original_filename, stored_filename, file_path, file_size, content_type, uploaded_by, image_width, image_height', 'Uploaded ultrasound images (raw bytes, no re-encode)'],
      ['24', 'ultrasound_annotations', 'id, order_id, image_id, annotation_source (AI|MANUAL), annotation_type, annotation_data (JSON), review_status, version, is_current, created_by, reviewed_by', 'Manual drawings/annotations on ultrasound images. Multiple versions supported.'],
      ['25', 'ultrasound_reports', 'id, test_order_id, version, image_description, professional_findings, conclusion, report_status (DRAFT|FINAL), signed_by_user_id, signed_name, signed_at, doctor_confirmed_by, doctor_confirmed_at', 'Final ultrasound reports with electronic signature.'],
      ['26', 'ai_analysis_results', 'id, test_order_id, status, detected, confidence, message, input_image, result_image, mask_image, xmin, ymin, xmax, ymax, error_message', 'AI analysis results with bounding box coordinates. Status: Success/Error/ManualConfirmed.'],
      ['27', 'audit_logs', 'id, user_id, action, table_name, old_value, new_value, ip_address, created_at', 'System audit trail for all important operations'],
      ['28', 'notifications', 'id, user_id, title, content, channel, is_read, created_at', 'In-app notifications for all system events'],
      ['29', 'pregnancies', 'id, patient_id, start_date, estimated_due_date, actual_delivery_date, pregnancy_status, fetus_count', 'Pregnancy tracking linked to patient'],
      ['30', 'reviews', 'id, appointment_id, rating, comment, created_at', 'Patient reviews/ratings for doctors'],
      ['31', 'password_reset_tokens', 'id, user_id, token, expires_at, is_used, created_at', 'Password reset tokens with expiry'],
      ['32', 'appointment_services', 'id, appointment_id, service_id, price', 'Services attached to appointment (price snapshot)'],
  ],
  [0.5, 2.5, 6, 8])

doc.add_page_break()

H('5.2 Code Packages', 3)
P('Package structure of the CAMS project:')
T(['Package', 'Description', 'Convention'],
  [
      ['com.clinic.controller', '60+ Servlets handling HTTP requests. One servlet per feature/page. Named: {Actor}{Feature}Servlet (e.g., PatientBookingServlet, StaffQueueServlet, DoctorAppointmentServlet).', 'doGet() for page loads, doPost() for form submissions. RequestDispatcher to JSP. Session-based auth check.'],
      ['com.clinic.service', '21 Service classes containing ALL business logic. Orchestrate DAOs, validate, manage transactions. Named: {Domain}Service (e.g., StaffReceptionService, UltrasoundOrderService, AppointmentStageService).', 'Stateless. Called by Servlets. Throw IllegalArgumentException for business errors. Use DatabaseConfig.getConnection() with try-with-resources.'],
      ['com.clinic.dao', '33 DAO classes. One per table (mostly). Raw JDBC queries. Named: {Table}DAO (e.g., AppointmentDAO, InvoiceDAO, MedicalRecordDAO).', 'Return Model objects or primitive types. No business logic. Parameterized queries only (SQL injection prevention).'],
      ['com.clinic.model', '35+ POJOs and enums. Plain data objects with getters/setters. Include: Appointment, MedicalRecord, Invoice, Prescription, User, Patient, Doctor, ExamStage enum, etc.', 'Match DB column names. Date fields use java.time (LocalDate, LocalDateTime, LocalTime). Money uses BigDecimal.'],
      ['com.clinic.filter', '4 Servlet Filters in chain: EncodingFilter→AuthenticationFilter→AuthorizationFilter→CsrfFilter.', '@WebFilter annotation. Chain order defined in web.xml.'],
      ['com.clinic.config', 'DatabaseConfig (connection pool), GoogleConfig (OAuth client), AuthorizationConfig (URL whitelist per role), SlotHoldExpiryListener.', 'Configuration constants and DB connection management.'],
      ['com.clinic.utils', 'AuditUtil (logging), NotificationHelper (notifications), BCryptUtil (password hash), EncryptionUtil (PII encryption), PaymentHelper (invoice helpers), StaffValidator (staff-specific validations).', 'Static utility methods. No state.'],
      ['web/views/', '65+ JSP files organized by actor folder: admin/, manager/, doctors/, staff/, sonographer/, patient/, auth/, common/, errors/, home/.', 'JSP + JSTL + EL. No Java scriptlets. Vietnamese UI text. Bootstrap CSS.'],
  ],
  [3, 8, 6])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# II. REQUIREMENT SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════
H('II. Requirement Specifications', 1)

# ── SECTION 1: AUTHENTICATION ──
H('1. Authentication', 2)

# UC-01
H('1.1 UC-01_Register', 3)
P('This screen allows an unauthenticated guest to create a new patient account by providing full name, email, phone number, password, and accepting terms & privacy policy. The system validates all inputs, ensures unique email, sends verification email with token, and creates account with Patient role and UNVERIFIED status.', bold=False)
P('Related use cases: UC-01_Register, UC-04_Verify Email')
P('')
P('SQL Commands:', bold=True)
P('1/ Check for duplicate registration (email, phone, username):')
P("SELECT COUNT(*) FROM users WHERE email = ? OR phone = ? OR username = ?;")
P('2/ Get Patient role ID:')
P("SELECT id FROM roles WHERE role_name = 'Patient';")
P('3/ Create user account (status=INACTIVE until email verified):')
P("INSERT INTO users (username, full_name, email, password_hash, phone, role_id, status, verification_token, is_verified, auth_provider, is_deleted, created_at)")
P("VALUES (?, ?, ?, ?, ?, ?, 'INACTIVE', ?, 0, 'LOCAL', 0, GETDATE());")
P('4/ Create linked patient profile:')
P("INSERT INTO patients (user_id, full_name, phone_number) VALUES (?, ?, ?);")
P('5/ Audit log:')
P("INSERT INTO audit_logs (user_id, action, table_name, new_value, ip_address, created_at) VALUES (?, 'REGISTER', 'users', ?, ?, GETDATE());")

# UC-02
H('1.2 UC-02_Login System', 3)
H('a. User Login', 4)
P('Allows a registered user to authenticate using username + password, or via Google Sign-In. Supports account lockout after 5 failed attempts.')
P('Related use cases: UC-02_Login System')
P('')
P('SQL Commands:', bold=True)
P('1/ Locate account by username or email:')
P("SELECT id, username, full_name, password_hash, email, phone, status, role_id, is_verified FROM users WHERE (username = ? OR email = ?) AND is_deleted = 0;")
P('2/ Check account status (ACTIVE required, INACTIVE → verify email, LOCKED → contact admin):')
P("-- status checked in Java from result set")
P('3/ On successful login — update session, clear failed attempts:')
P("-- Session attributes set: userId, username, role, fullName")
P("INSERT INTO audit_logs (user_id, action, table_name, ip_address, created_at) VALUES (?, 'LOGIN_SUCCESS', 'users', ?, GETDATE());")
P('4/ On failed login — increment counter (handled in-memory or via DB):')
P("INSERT INTO audit_logs (user_id, action, table_name, ip_address, created_at) VALUES (?, 'LOGIN_FAILED', 'users', ?, GETDATE());")

H('b. Google Login — Auto-Provisioning', 4)
P('When user signs in with Google using an email not yet in the system, the system prompts for missing required fields (phone, name confirmation) before auto-creating a Patient account.')
P('')
P('SQL Commands:', bold=True)
P('1/ Verify Google email is not already registered:')
P("SELECT COUNT(*) FROM users WHERE email = ? AND is_deleted = 0;")
P('2/ Auto-create account (status=ACTIVE, is_verified=1, auth_provider=GOOGLE):')
P("INSERT INTO users (username, full_name, email, role_id, status, is_verified, google_id, auth_provider, is_deleted, created_at)")
P("VALUES (?, ?, ?, (SELECT id FROM roles WHERE role_name='Patient'), 'ACTIVE', 1, ?, 'GOOGLE', 0, GETDATE());")
P('3/ Create linked patient profile:')
P("INSERT INTO patients (user_id, full_name) VALUES (?, ?);")

# UC-03
H('1.3 UC-03_Forgot Password', 3)
P('Allows a registered user to reset a forgotten password. User enters email → receives OTP/reset link → sets new password.')
P('')
P('SQL Commands:', bold=True)
P('1/ Verify email exists:')
P("SELECT id, full_name, email FROM users WHERE email = ? AND is_deleted = 0;")
P('2/ Generate reset token (valid for 1 hour):')
P("INSERT INTO password_reset_tokens (user_id, token, expires_at, is_used, created_at) VALUES (?, ?, DATEADD(HOUR, 1, GETDATE()), 0, GETDATE());")
P('3/ Reset password (token must be valid, unexpired, unused):')
P("UPDATE users SET password_hash = ?, updated_at = GETDATE() WHERE id = ?;")
P("UPDATE password_reset_tokens SET is_used = 1 WHERE id = ?;")

# UC-04
H('1.4 UC-04_Change Password', 3)
P('Allows logged-in user to change password. Requires current password verification.')
P('')
P('SQL Commands:', bold=True)
P('1/ Verify current password (BCrypt check in Java against stored hash):')
P("SELECT password_hash FROM users WHERE id = ?;")
P('2/ Update to new password:')
P("UPDATE users SET password_hash = ?, updated_at = GETDATE() WHERE id = ?;")

# UC-05
H('1.5 UC-05_Logout', 3)
P('Terminates user session. Clears all session attributes. Redirects to login page.')
P('')
P('SQL Commands:', bold=True)
P("INSERT INTO audit_logs (user_id, action, table_name, ip_address, created_at) VALUES (?, 'LOGOUT', 'users', ?, GETDATE());")

doc.add_page_break()

# ── SECTION 2: APPOINTMENT BOOKING ──
H('2. Appointment Booking', 2)

H('2.1 UC-06_Book Appointment Online (Patient)', 3)
P('Patient books appointment: selects doctor → system displays exam fee → picks date → available slots shown → chooses slot → enters symptoms (10-500 chars mandatory) and LMP (if female) → confirms. System shows total with "Thanh toán tại quầy lễ tân trước giờ khám". On confirm: creates Pending appointment with base_fee locked.')
P('Related use cases: UC-06_Book Appointment, UC-07_View Available Slots')
P('')
P('Business Rules (all checked at SERVER):', bold=True)
B('Slot must exist in doctor_schedules with status=APPROVED and booked_count < max_slots')
B('Date must not be in the past')
B('Patient must not have another active appointment (Pending/Confirmed/Waiting/InProgress) on the same day')
B('Symptoms: 10-500 characters, at least 2 words, not numbers only, no repeating chars >5x')
B('LMP: not after appointment date, not in future, gestational age ≤42 weeks (294 days)')
B('Phone: starts with 0, 10-11 digits')
B('Age: 10-100 years from date_of_birth')
B('Anti-duplicate: double-click/F5 only creates 1 appointment')
B('base_fee = 200,000đ + (doctor.experience_years × 50,000đ) — calculated once and stored')
P('')
P('SQL Commands:', bold=True)
P('1/ Get doctor with experience_years for fee calculation:')
P("SELECT d.* FROM doctors d WHERE d.id = ?;")
P('2/ Check slot availability:')
P("SELECT ds.id, ds.booked_count, ds.max_slots, ds.status, ds.work_date FROM doctor_schedules ds WHERE ds.id = ? AND ds.status = 'APPROVED';")
P('3/ Check duplicate same-day appointment:')
P("SELECT COUNT(*) FROM appointments WHERE patient_id = ? AND appointment_date = ? AND status IN ('Pending','Confirmed','Waiting','InProgress');")
P('4/ Create appointment with locked base_fee (in transaction with slot hold):')
P("INSERT INTO appointments (patient_id, doctor_id, appointment_date, booking_source, symptoms, last_menstrual_period, status, time_slot, base_fee, schedule_id, is_priority)")
P("VALUES (?, ?, ?, 'Online', ?, ?, 'Pending', ?, ?, ?, 0);")
P('5/ Increment booked_count:')
P("UPDATE doctor_schedules SET booked_count = booked_count + 1 WHERE id = ? AND booked_count < max_slots;")
P('6/ Notify doctor:')
P("INSERT INTO notifications (user_id, title, content, channel, is_read, created_at) VALUES (?, N'Lịch hẹn mới', ?, 'System', 0, GETDATE());")

H('2.2 UC-07_View Available Slots (API)', 3)
P('AJAX endpoint returning JSON list of available schedule slots for a doctor on a given date. Only APPROVED schedules with available capacity.')
P('')
P('SQL Commands:', bold=True)
P("SELECT ds.id, ds.doctor_id, ds.work_date, ds.max_slots, ds.booked_count, s.name AS shift_name, s.start_time, s.end_time")
P("FROM doctor_schedules ds JOIN shifts s ON ds.shift_id = s.id")
P("WHERE ds.doctor_id = ? AND ds.work_date = ? AND ds.status = 'APPROVED' AND ds.booked_count < ds.max_slots")
P("ORDER BY s.start_time;")

H('2.3 UC-08_Create Manual Appointment (Staff)', 3)
P('Staff creates appointment for walk-in patient. Looks up patient by phone (or creates new). Selects doctor, date, slot. Enters symptoms. System auto-creates Confirmed appointment + PRE_EXAM invoice in one flow. Option: immediate check-in if today and override reason provided.')
P('')
P('SQL Commands:', bold=True)
P('1/ Find patient by phone:')
P("SELECT p.* FROM patients p WHERE p.phone_number = ?;")
P('2/ Create patient if not found:')
P("INSERT INTO patients (full_name, phone_number, date_of_birth) VALUES (?, ?, ?);")
P('3/ Create appointment + hold slot + base_fee (transactional — StaffReceptionService.createManualBookingWithOverride()):')
P("INSERT INTO appointments (...) VALUES (...); UPDATE doctor_schedules SET booked_count = booked_count + 1 WHERE id = ?;")
P('4/ Create PRE_EXAM invoice (copy base_fee from appointment):')
P("INSERT INTO invoices (appointment_id, total_amount, status, invoice_type, created_at) VALUES (?, ?, 'Unpaid', 'PRE_EXAM', GETDATE());")

doc.add_page_break()

# ── SECTION 3: RECEPTION & QUEUE ──
H('3. Reception & Queue Management', 2)

H('3.1 UC-09_View Reception Queue', 3)
P('Staff views paginated list of today\'s appointments. Filterable by status and search (name/phone/appointment ID). Sorted by: status priority (Pending first → InProgress → Waiting → Confirmed → Completed → Cancelled → NoShow last), then priority flag, then created_at.')
P('')
P('SQL Commands:', bold=True)
P("SELECT a.*, p.full_name AS patient_name, p.phone_number, d.full_name AS doctor_name")
P("FROM appointments a LEFT JOIN patients p ON a.patient_id = p.id LEFT JOIN doctors d ON a.doctor_id = d.id")
P("WHERE a.appointment_date = ? AND a.status NOT IN ('Cancelled','NoShow')")
P("ORDER BY CASE a.status WHEN 'Pending' THEN 1 WHEN 'InProgress' THEN 2 WHEN 'Waiting' THEN 3 WHEN 'Confirmed' THEN 4 ELSE 9 END, a.is_priority DESC, a.created_at ASC;")

H('3.2 UC-10_Approve Booking', 3)
P('Staff approves a Pending appointment. System changes status to Confirmed and creates PRE_EXAM invoice (Unpaid) with amount copied from appointment.base_fee. No recalculation.')
P('')
P('Business Rules:', bold=True)
B('Appointment must be in Pending status')
B('PRE_EXAM must not already exist (dedup check)')
B('Invoice amount = appointment.base_fee (locked price, never recalculated)')
B('Patient receives notification: "Vui lòng đến quầy lễ tân thanh toán trước giờ khám" with amount')
P('')
P('SQL Commands:', bold=True)
P('1/ Update appointment status:')
P("UPDATE appointments SET status = 'Confirmed' WHERE id = ? AND status = 'Pending';")
P('2/ Create PRE_EXAM invoice:')
P("INSERT INTO invoices (appointment_id, total_amount, status, transaction_code, invoice_type, created_at)")
P("SELECT ?, base_fee, 'Unpaid', ?, 'PRE_EXAM', GETDATE() FROM appointments WHERE id = ?")
P("WHERE NOT EXISTS (SELECT 1 FROM invoices WHERE appointment_id = ? AND invoice_type = 'PRE_EXAM');")
P('3/ Send notification + audit log.')

H('3.3 UC-11_Confirm Payment (PRE_EXAM)', 3)
P('Staff confirms patient has paid exam fee in cash at counter. Updates invoice to Paid with full payment evidence: paid_at, paid_by_user_id, payment_method=CASH.')
P('')
P('Business Rules:', bold=True)
B('Staff role required')
B('Appointment status: Confirmed or Waiting')
B('PRE_EXAM invoice must exist and be Unpaid')
B('Anti-duplicate: "Hoá đơn PRE_EXAM đã được thanh toán trước đó. Không thể xác nhận trùng."')
B('Records: paid_at = GETDATE(), paid_by_user_id = staff_user_id, payment_method = CASH')
B('If invoice doesn\'t exist yet (edge case), INSERT new Paid invoice')
P('')
P('SQL Commands:', bold=True)
P('1/ Update invoice to Paid:')
P("UPDATE invoices SET status = 'Paid', transaction_code = ?, confirmed_at = GETDATE(), confirmed_by = ?, paid_at = GETDATE(), paid_by_user_id = ?, payment_method = 'CASH'")
P("WHERE appointment_id = ? AND invoice_type = 'PRE_EXAM' AND status = 'Unpaid';")
P('2/ Insert if not exists (UPSERT pattern with NOT EXISTS subquery):')
P("INSERT INTO invoices (appointment_id, invoice_type, total_amount, status, transaction_code, confirmed_at, confirmed_by, paid_at, paid_by_user_id, payment_method)")
P("SELECT ?, 'PRE_EXAM', ?, 'Paid', ?, GETDATE(), ?, GETDATE(), ?, 'CASH'")
P("WHERE NOT EXISTS (SELECT 1 FROM invoices WHERE appointment_id = ? AND invoice_type = 'PRE_EXAM');")

H('3.4 UC-12_Confirm Payment (POST_EXAM)', 3)
P('Staff confirms payment for ultrasound service invoice. Similar to UC-11 but for POST_EXAM type. Additionally backfills invoice_items if missing.')
P('')
P('Business Rules:', bold=True)
B('Appointment must be InProgress')
B('POST_EXAM must be Unpaid')
B('On payment: backfill invoice_items from test_orders.service_id + price')

H('3.5 UC-13_Check In Patient', 3)
P('Staff checks in patient. Changes status from Confirmed to Waiting, assigns queue_number. All preconditions checked at server.')
P('')
P('Business Rules (ALL MANDATORY, checked sequentially):', bold=True)
B('1. Appointment exists')
B('2. appointment_date = TODAY')
B('3. Not Cancelled, not SUCCESS/Completed, not InProgress, not Waiting (already checked in)')
B('4. Not Pending → "Lịch hẹn chưa được duyệt. Vui lòng nhấn Duyệt & Tạo Hóa Đơn trước."')
B('5. PRE_EXAM invoice exists → if not: "Chưa có hóa đơn phí khám."')
B('6. PRE_EXAM.status = Paid → if not: "Bệnh nhân chưa thanh toán phí khám."')
B('7. Within check-in window: ≤120 minutes before shift start, before shift end')
P('')
P('SQL Commands:', bold=True)
P('1/ Check PRE_EXAM paid:')
P("SELECT status FROM invoices WHERE appointment_id = ? AND invoice_type = 'PRE_EXAM';")
P('2/ Check shift time:')
P("SELECT s.start_time, s.end_time FROM doctor_schedules ds JOIN shifts s ON ds.shift_id = s.id WHERE ds.id = ?;")
P('3/ Check-in and renumber queue:')
P("UPDATE appointments SET status = 'Waiting', queue_number = ? WHERE id = ? AND status = 'Confirmed';")
P('-- queue_number assignment takes into account is_priority flag and existing queue')

H('3.6 UC-14_Mark Priority', 3)
P('Staff marks a waiting patient as priority. Requires reason (5-500 chars). Renumbers the entire queue for the day.')
P('')
P('Business Rules:', bold=True)
B('Appointment status = Waiting')
B('appointment_date = TODAY')
B('PRE_EXAM must be Paid')
B('Not already priority')
B('Reason: 5-500 characters')
P('')
P('SQL Commands:', bold=True)
P("UPDATE appointments SET is_priority = 1, priority_reason = ?, prioritized_at = GETDATE(), prioritized_by = ? WHERE id = ?;")
P("-- Then renumber all Waiting appointments for today by (is_priority DESC, created_at ASC)")

H('3.7 UC-15_Cancel Appointment (Staff)', 3)
P('Staff cancels appointment. Releases slot. If PRE_EXAM was paid, auto-refunds with alert for staff to handle cash with patient.')
P('')
P('Business Rules:', bold=True)
B('Not already Cancelled, not SUCCESS/Completed, not InProgress, not Waiting, not NoShow')
B('If PRE_EXAM Paid → refund first, then cancel')
B('Transaction: update status + release slot (booked_count - 1) + cancel unpaid invoices')
P('')
P('SQL Commands:', bold=True)
P('1/ Cancel appointment + release slot:')
P("BEGIN TRANSACTION;")
P("UPDATE appointments SET status = 'Cancelled' WHERE id = ?;")
P("UPDATE doctor_schedules SET booked_count = CASE WHEN booked_count > 0 THEN booked_count - 1 ELSE 0 END WHERE id = ?;")
P("UPDATE invoices SET status = 'Cancelled' WHERE appointment_id = ? AND status = 'Unpaid';")
P("COMMIT;")

H('3.8 UC-16_Mark No Show', 3)
P('Staff marks patient as no-show. Only for today\'s Pending/Confirmed appointments after shift has ended.')
P('')
P('Business Rules:', bold=True)
B('Status: Pending or Confirmed')
B('appointment_date = TODAY')
B('Shift has ended (current time > shift end_time)')
B('If PRE_EXAM was Paid → warning (staff must handle refund separately)')
P('')
P('SQL Commands:', bold=True)
P("UPDATE appointments SET status = 'NoShow' WHERE id = ?;")
P("UPDATE doctor_schedules SET booked_count = CASE WHEN booked_count > 0 THEN booked_count - 1 ELSE 0 END WHERE id = ?;")
P("UPDATE invoices SET status = 'Cancelled' WHERE appointment_id = ? AND status = 'Unpaid';")

H('3.9 UC-17_Refund Invoice', 3)
P('Staff refunds a paid invoice. Changes status to Refunded. Records refunded_at, refunded_by_user_id, refund_reason. Revenue reports subtract refunded amounts.')
P('')
P('Business Rules:', bold=True)
B('Invoice status must be Paid')
B('Cannot refund twice (check status=Paid in WHERE clause)')
B('Reason: 10-500 characters')
B('If POST_EXAM: auto-cancel linked test_order')
B('Revenue = SUM(Paid) - SUM(Refunded) — Manager sees real-time adjustment')
P('')
P('SQL Commands:', bold=True)
P("UPDATE invoices SET status = 'Refunded', refunded_at = GETDATE(), refunded_by_user_id = ?, refund_reason = ? WHERE id = ? AND status = 'Paid';")
P("UPDATE test_orders SET status = 'Cancelled' WHERE id IN (SELECT o.id FROM test_orders o JOIN medical_records mr ON o.medical_record_id = mr.id WHERE mr.appointment_id = ?) AND status IN ('Pending','Ordered');")

doc.add_page_break()

# ── SECTIONS 4-15: Continue with remaining use cases ──
# (Abbreviated for length; same detailed format)

H('4. Clinical Examination', 2)
P('See detailed Design Specifications in Part III, Section 4.', bold=True)

H('5. Ultrasound Diagnostics & AI Integration', 2)
P('See detailed Design Specifications in Part III, Section 5.', bold=True)

H('6. Diagnosis & Prescription', 2)
P('See detailed Design Specifications in Part III, Section 6.', bold=True)

H('7. Payment Management', 2)
P('See detailed Design Specifications in Part III, Section 7.', bold=True)

H('8. Medical Records', 2)
P('See detailed Design Specifications in Part III, Section 4.', bold=True)

H('9. Patient Portal', 2)
P('See detailed Design Specifications in Part III, Section 9.', bold=True)

H('10. Service & Price Management', 2)
P('See detailed Design Specifications in Part III, Section 9.', bold=True)

H('11. Doctor Schedule Management', 2)
P('See detailed Design Specifications in Part III, Section 10.', bold=True)

H('12. Revenue & Statistics', 2)
P('See detailed Design Specifications in Part III, Section 8.', bold=True)

H('13. User & Role Management', 2)
P('See detailed Design Specifications in Part III, Section 10.', bold=True)

H('14. Pharmacy Management', 2)
P('See detailed Design Specifications in Part III, Section 10.', bold=True)

H('15. System & Security', 2)
P('See detailed Design Specifications in Part III, Section 10.', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# III. DESIGN SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════
H('III. Design Specifications', 1)

# ── 1. AUTHENTICATION ──
H('1. Authentication', 2)

H('1.1 UC-01_Register', 3)
P('This screen allows an unauthenticated guest to create a new patient account by providing personal information. The system validates all inputs, ensures unique username/email/phone, and creates a new account with Patient role, INACTIVE status, pending email verification.')
P('Related use cases: UC-01_Register')
P('')
P('UI Design:', bold=True)
B('Form fields: Full Name*, Email*, Phone*, Username*, Password*, Confirm Password*, Terms checkbox*')
B('Client-side: JS validation for required fields, password match, email format, phone format')
B('Server-side: ALL validations repeated (client is cosmetic only per CAMS rules)')
B('Success: "Vui lòng kiểm tra email để xác nhận tài khoản" → redirect login')
B('Error: Vietnamese message for each validation failure')
P('')
P('Database Access — SQL Commands:', bold=True)
P('1/ Check for duplicate registration information (Username, Email, Phone):')
P("SELECT COUNT(*) FROM users WHERE username = ?;")
P("SELECT COUNT(*) FROM users WHERE email = ?;")
P("SELECT COUNT(*) FROM users WHERE phone = ?;")
P('2/ Get the Patient role ID:')
P("SELECT id FROM roles WHERE role_name = 'Patient';")
P('3/ Create new user with INACTIVE status and verification token:')
P("INSERT INTO users (username, full_name, email, password_hash, phone, role_id, status, verification_token, is_verified, auth_provider, is_deleted, created_at)")
P("VALUES (?, ?, ?, ?, ?, ?, 'INACTIVE', ?, 0, 'LOCAL', 0, GETDATE());")
P('4/ Create linked patient profile:')
P("INSERT INTO patients (user_id, full_name, phone_number) VALUES (?, ?, ?);")
P('5/ Audit log:')
P("INSERT INTO audit_logs (user_id, action, table_name, new_value, ip_address, created_at) VALUES (?, 'REGISTER', 'users', ?, ?, GETDATE());")

H('1.2 UC-02_Login System', 3)
H('a. User Login', 4)
P('This screen allows a registered user to authenticate into the system using username + password, or via Google Sign-In.')
P('Related use cases: UC-02_Login System')
P('')
P('UI Design:', bold=True)
B('Form: Username/Email field + Password field + Login button + Google Sign-In button')
B('Links: Forgot Password, Register')
B('Error messages in Vietnamese: "Tên đăng nhập hoặc mật khẩu không đúng", "Tài khoản chưa được xác nhận email"')
B('On success: redirect to role-based home dashboard (/patient/home, /doctor/dashboard, etc.)')
P('')
P('Database Access — SQL Commands:', bold=True)
P('1/ Locate account by username or email:')
P("SELECT id, username, full_name, password_hash, email, phone, status, role_id, is_verified FROM users WHERE (username = ? OR email = ?) AND is_deleted = 0;")
P('2/ Verify password with BCrypt (in Java):')
P('-- BCrypt.checkpw(plainPassword, storedHash)')
P('3/ Check account status: ACTIVE → allow; INACTIVE → "Vui lòng xác nhận email"; LOCKED → "Tài khoản bị khoá"')
P('4/ On success — create session, log audit:')
P("INSERT INTO audit_logs (user_id, action, table_name, ip_address, created_at) VALUES (?, 'LOGIN_SUCCESS', 'users', ?, GETDATE());")

H('b. Google Login — Complete Registration (Auto-Provisioning)', 4)
P('When a user signs in with Google using an email that does not yet exist in the system, the system prompts for the remaining required fields before creating a new PATIENT account automatically.')
P('')
P('SQL Commands:', bold=True)
P('1/ Verify the Google email is not already registered:')
P("SELECT COUNT(*) FROM users WHERE email = ? AND is_deleted = 0;")
P('2/ Get the PATIENT role:')
P("SELECT id FROM roles WHERE role_name = 'Patient';")
P('3/ Create the new account (ACTIVE, verified, GOOGLE provider):')
P("INSERT INTO users (username, full_name, email, role_id, status, is_verified, google_id, auth_provider, is_deleted, created_at)")
P("VALUES (?, ?, ?, ?, 'ACTIVE', 1, ?, 'GOOGLE', 0, GETDATE());")
P('4/ Create linked Patient profile:')
P("INSERT INTO patients (user_id, full_name) VALUES (?, ?);")

H('1.3 UC-03_Forgot Password', 3)
P('Allows a registered user to reset a forgotten password through a 3-step process: Enter Email → Receive OTP/Token → Set New Password.')
P('')
P('Step 1 — Enter Email:', bold=True)
P('SQL: SELECT id, full_name, email FROM users WHERE email = ? AND is_deleted = 0;')
P('Step 2 — Generate Reset Token (1-hour expiry):')
P("INSERT INTO password_reset_tokens (user_id, token, expires_at, is_used, created_at) VALUES (?, ?, DATEADD(HOUR, 1, GETDATE()), 0, GETDATE());")
P('Step 3 — Reset Password (token validated in Java):')
P("UPDATE users SET password_hash = ?, updated_at = GETDATE() WHERE id = ?;")
P("UPDATE password_reset_tokens SET is_used = 1 WHERE id = ?;")

H('1.4 UC-04_Change Password', 3)
P('Logged-in user changes password. Requires current password confirmation + new password (with strength rules).')
P('')
P('SQL Commands:', bold=True)
P("SELECT password_hash FROM users WHERE id = ?;  -- verify current password with BCrypt")
P("UPDATE users SET password_hash = ?, updated_at = GETDATE() WHERE id = ?;  -- set new password")

doc.add_page_break()

# ── 2. APPOINTMENT BOOKING ──
H('2. Appointment Booking', 2)

H('2.1 UC-06_Book Appointment Online', 3)
P('Multi-step booking form for Patient. Step 1: Select doctor (card view with name, specialty, experience, price). Step 2: Select date. Step 3: Select available time slot from approved doctor_schedules. Step 4: Enter symptoms (10-500 chars), LMP (if applicable). Step 5: Confirm — shows all info + total price + "Thanh toán tại quầy lễ tân trước giờ khám".')
P('')
P('UI Design:', bold=True)
B('Doctor cards showing: avatar, full_name, specialization, experience_years, calculated exam fee')
B('Date picker (min: today, blocked: dates with no approved schedules for selected doctor)')
B('Slot list: radio buttons showing shift name + time range + remaining slots')
B('Form: symptoms textarea with char counter, LMP date picker (conditional on gender)')
B('Confirmation modal: summary table with doctor name, date, time, symptoms, total fee')
B('Success: redirect to appointment list with green success message')
P('')
P('Database Access — SQL Commands:', bold=True)
P('1/ Get all active doctors with experience_years (for fee display):')
P("SELECT d.*, u.full_name FROM doctors d JOIN users u ON d.user_id = u.id JOIN roles r ON u.role_id = r.id WHERE r.role_name = 'Doctor' AND u.status = 'ACTIVE';")
P('2/ Get approved schedules with available slots for a doctor on a date:')
P("SELECT ds.id, ds.doctor_id, ds.work_date, ds.max_slots, ds.booked_count, s.name AS shift_name, s.start_time, s.end_time")
P("FROM doctor_schedules ds JOIN shifts s ON ds.shift_id = s.id")
P("WHERE ds.doctor_id = ? AND ds.work_date = ? AND ds.status = 'APPROVED' AND ds.booked_count < ds.max_slots;")
P('3/ Check duplicate same-day appointment:')
P("SELECT COUNT(*) FROM appointments WHERE patient_id = ? AND appointment_date = ? AND status IN ('Pending','Confirmed','Waiting','InProgress');")
P('4/ Create appointment with locked base_fee (transactional with slot increment):')
P("BEGIN TRANSACTION;")
P("INSERT INTO appointments (patient_id, doctor_id, appointment_date, booking_source, symptoms, last_menstrual_period, status, time_slot, base_fee, schedule_id, is_priority, held_by, held_until)")
P("VALUES (?, ?, ?, 'Online', ?, ?, 'Pending', ?, ?, ?, 0, ?, DATEADD(MINUTE, 5, GETDATE()));")
P("UPDATE doctor_schedules SET booked_count = booked_count + 1 WHERE id = ? AND booked_count < max_slots;")
P("COMMIT;")
P('5/ Notify doctor:')
P("SELECT user_id FROM doctors WHERE id = ?;")
P("INSERT INTO notifications (user_id, title, content, channel, is_read, created_at) VALUES (?, N'Lịch hẹn mới', ?, 'System', 0, GETDATE());")

doc.add_page_break()

# ── 3. RECEPTION & QUEUE MANAGEMENT ──
H('3. Reception & Queue Management', 2)

H('3.1 UC-09_View Reception Queue', 3)
P('Staff\'s main working screen. Shows all today\'s appointments as a paginated, filterable, searchable table. Each row shows: appointment ID, patient name, phone, doctor name, time slot, status (Vietnamese), queue number, priority flag. Action buttons appear conditionally based on appointment status and business rules.')
P('')
P('UI Design:', bold=True)
B('Top bar: date selector (default today), status filter dropdown, search box (name/phone/ID)')
B('Table columns: ID, Patient, Phone, Doctor, Time, Status (colored badge), Queue#, Priority (flag icon), Actions')
B('Color-coded status badges: Pending=yellow, Confirmed=blue, Waiting=green, InProgress=orange, Success=green-bold, Cancelled=red, NoShow=gray')
B('Conditional action buttons per row based on status + payment state')
B('Late indicator: red highlight for Confirmed patients past check-in deadline')
B('Auto-refresh: page reloads on action; autoMovePatientsToNextShift runs on each page load')
P('')
P('Database Access — SQL Commands:', bold=True)
P('1/ Get paginated appointments for today:')
P("SELECT a.*, p.full_name AS patient_name, p.phone_number, d.full_name AS doctor_name, s.name AS shift_name, s.start_time, s.end_time")
P("FROM appointments a LEFT JOIN patients p ON a.patient_id = p.id LEFT JOIN doctors d ON a.doctor_id = d.id")
P("LEFT JOIN doctor_schedules ds ON a.schedule_id = ds.id LEFT JOIN shifts s ON ds.shift_id = s.id")
P("WHERE a.appointment_date = ? AND a.status NOT IN ('Cancelled','NoShow')")
P("ORDER BY CASE a.status WHEN 'Pending' THEN 1 WHEN 'InProgress' THEN 2 WHEN 'Waiting' THEN 3 WHEN 'Confirmed' THEN 4 ELSE 9 END, a.is_priority DESC, a.created_at ASC")
P("OFFSET ? ROWS FETCH NEXT ? ROWS ONLY;")
P('2/ Check PRE_EXAM payment status per appointment (shown in UI for action eligibility):')
P("SELECT status FROM invoices WHERE appointment_id = ? AND invoice_type = 'PRE_EXAM';")
P('3/ Get doctor workload for today (for UI info):')
P("SELECT doctor_id, COUNT(*) AS cnt FROM appointments WHERE appointment_date = ? AND status NOT IN ('Cancelled','NoShow','SUCCESS','Completed') GROUP BY doctor_id;")

H('3.2 UC-10 to UC-17 — Staff Actions', 3)
P('All staff actions (Approve, Confirm Payment, Check-in, Mark Priority, Cancel, NoShow, Refund) are POST endpoints handled by StaffQueueServlet. Detailed SQL and validation are documented in Part II Sections 3.2-3.9 above. Key design principle: EVERY action validates preconditions at SERVER, uses transactions for multi-table writes, and writes audit logs.')

doc.add_page_break()

# ── 4. CLINICAL EXAMINATION ──
H('4. Clinical Examination & Medical Record', 2)

H('4.1 UC-18_View Today\'s Appointments & Accept Patient', 3)
P('Doctor views list of today\'s patients assigned to them. Each row shows: queue number, patient name, status, time. Doctor clicks "Tiếp nhận" to accept a Waiting patient → status changes to InProgress.')
P('')
P('SQL Commands:', bold=True)
P('1/ Get today\'s appointments for this doctor:')
P("SELECT a.*, p.full_name AS patient_name, p.date_of_birth, p.phone_number, a.queue_number")
P("FROM appointments a JOIN patients p ON a.patient_id = p.id WHERE a.doctor_id = ? AND a.appointment_date = ? AND a.status != 'Cancelled' ORDER BY a.queue_number;")
P('2/ Accept patient (Waiting → InProgress):')
P("UPDATE appointments SET status = 'InProgress' WHERE id = ? AND doctor_id = ? AND status = 'Waiting' AND appointment_date = ?;")

H('4.2 UC-19_Update Medical Record (Clinical Exam)', 3)
P('Doctor fills the clinical examination form — Block 2 of the 5-block accordion. Records vitals, symptoms, obstetric history, risk factors. Saves as Draft medical_record.')
P('')
P('UI Design (5-block accordion):', bold=True)
B('Block 1 (Accept): Single button "Tiếp nhận ca khám". Collapses after acceptance with checkmark + "Đã tiếp nhận".')
B('Block 2 (Clinical Exam): Form with: weight, height, blood pressure, pulse, temperature, gestational age (weeks+days), fundal height, fetal heart rate, fetal presentation/position/movement, cervical dilation/effacement, amniotic fluid, edema, proteinuria, vaginal bleeding, uterine contractions, clinical notes, risk flags. Save button "Lưu nháp".')
B('Block 3 (Imaging Decision): Radio: "Không chỉ định" or "Chỉ định siêu âm". If ultrasound: dropdown of active services + reason textarea (10-500 chars mandatory). Cancel button (conditional).')
B('Block 4 (Diagnosis): Read-only ultrasound results panel (images, AI result, sonographer conclusion). "Xác nhận đã xem" button. Diagnosis textarea, treatment plan, prescription sub-form. Purchase decision: AT_CLINIC or EXTERNAL.')
B('Block 5 (Finalize): "Chốt hồ sơ bệnh án" button. After finalization: entire page read-only.')
P('')
P('Lock/Unlock logic — getStage() single source of truth:', bold=True)
B('CLINICAL_EXAM: Block 2 open, Blocks 3-5 locked with reason "Cần lưu bệnh án nháp trước"')
B('ORDER_DECISION: Block 3 open, Blocks 4-5 locked with reason "Vui lòng quyết định chỉ định cận lâm sàng"')
B('WAITING_PAYMENT: Blocks 4-5 locked with reason "Đang chờ bệnh nhân thanh toán dịch vụ siêu âm"')
B('WAITING_ULTRASOUND: Blocks 4-5 locked with reason "Đang chờ bác sĩ siêu âm trả kết quả" + refresh button')
B('DIAGNOSIS: Block 4 open, Block 5 locked with reason "Cần nhập chẩn đoán và quyết định đơn thuốc"')
B('READY_TO_FINALIZE: Block 5 open with finalize button')
B('FINALIZED: All blocks collapsed, read-only, green banner "Hồ sơ đã hoàn tất"')
P('')
P('Database Access — SQL Commands:', bold=True)
P('1/ Load medical record for appointment (creates if not exists on first save):')
P("SELECT * FROM medical_records WHERE appointment_id = ?;")
P('2/ Upsert clinical exam data:')
P("MERGE medical_records AS target USING (SELECT ? AS appointment_id) AS source ON target.appointment_id = source.appointment_id")
P("WHEN MATCHED THEN UPDATE SET clinical_notes = ?, weight_kg = ?, blood_pressure = ?, pulse_bpm = ?, temperature_c = ?, height_cm = ?, gestational_age_weeks = ?, gestational_age_days = ?, fundal_height_cm = ?, fetal_heart_rate = ?, fetal_presentation = ?, fetal_position = ?, fetal_movement = ?, cervical_dilation_cm = ?, cervical_effacement = ?, amniotic_fluid = ?, edema = ?, proteinuria = ?, vaginal_bleeding = ?, uterine_contractions = ?, risk_flags_json = ?, updated_at = GETDATE(), updated_by = ?")
P("WHEN NOT MATCHED THEN INSERT (appointment_id, clinical_notes, weight_kg, ..., status) VALUES (?, ?, ?, ..., 'Draft');")
P('3/ Get all prescriptions for medical record:')
P("SELECT rx.*, pi.id AS item_id, pi.medicine_id, pi.quantity, pi.dosage, m.name AS med_name")
P("FROM prescriptions rx LEFT JOIN prescription_items pi ON rx.id = pi.prescription_id LEFT JOIN medicines m ON pi.medicine_id = m.id")
P("WHERE rx.medical_record_id = ?;")
P('4/ Get ultrasound results for display in Block 4 (read-only):')
P("SELECT tor.id AS order_id, tor.status AS order_status, s.service_name, ur.image_description, ur.professional_findings, ur.conclusion, ur.signed_name, ur.signed_at")
P("FROM test_orders tor LEFT JOIN services s ON tor.service_id = s.id LEFT JOIN ultrasound_reports ur ON tor.id = ur.test_order_id AND ur.is_current = 1")
P("WHERE tor.medical_record_id = ? AND tor.status != 'Cancelled' ORDER BY tor.id;")

doc.add_page_break()

# ── 5. ULTRASOUND DIAGNOSTICS ──
H('5. Ultrasound Diagnostics & AI Integration', 2)

H('5.1 UC-23_View Ultrasound Queue', 3)
P('Sonographer dashboard showing all ultrasound orders. Filterable by status and emergency flag. Shows: order ID, patient name, service name, doctor name, ordered date, status, emergency indicator.')
P('')
P('SQL Commands:', bold=True)
P('1/ Get paginated ultrasound orders (ready for sonographer or InProgress):')
P("SELECT tor.id, tor.status, tor.created_at, tor.service_id, s.service_name, mr.id AS mr_id, a.id AS appt_id, a.appointment_date, p.full_name AS patient_name, d.full_name AS doctor_name")
P("FROM test_orders tor JOIN medical_records mr ON tor.medical_record_id = mr.id")
P("JOIN appointments a ON mr.appointment_id = a.id JOIN patients p ON a.patient_id = p.id")
P("JOIN doctors d ON tor.doctor_id = d.id JOIN services s ON tor.service_id = s.id")
P("WHERE tor.status IN ('Ordered','InProgress') ORDER BY a.is_emergency DESC, tor.created_at ASC;")

H('5.2 UC-24_Perform Ultrasound (4-Step Flow)', 3)
P('Step-by-step ultrasound examination with mandatory sequential flow enforcement at server.')
P('')
P('Step 1 — Accept Case:', bold=True)
P('Sonographer clicks "Tiếp nhận ca". System validates: POST_EXAM paid, no other sonographer assigned. Sets sonographer_user_id + accepted_at.')
P("UPDATE test_orders SET sonographer_user_id = ?, accepted_at = GETDATE(), status = 'InProgress' WHERE id = ? AND sonographer_user_id IS NULL;")
P('')
P('Step 2 — Upload Images:', bold=True)
P('Sonographer uploads ultrasound images. Stored as raw bytes (NO ImageIO re-encode per CAMS rules). Saved to /uploads/ultrasound/ with UUID filename.')
P("INSERT INTO ultrasound_images (test_order_id, original_filename, stored_filename, file_path, file_size, content_type, uploaded_by, image_width, image_height, uploaded_at)")
P("VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, GETDATE());")
P('')
P('Step 3 — AI Analysis:', bold=True)
P('System sends image to Python AI Engine via HTTP (AiPredictionService). Timeout configured. On success: bounding box + confidence + message stored. On failure: "AI Engine không khả dụng. Bạn có thể vẽ thủ công." — allows skip.')
P("INSERT INTO ai_analysis_results (test_order_id, status, detected, confidence, message, input_image, result_image, mask_image, xmin, ymin, xmax, ymax, analyzed_at)")
P("VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, GETDATE());")
P('')
P('Two branches for Step 3:', bold=True)
B('A) Accept AI Result: Sonographer agrees with AI. Fills findings → proceeds to Sign.')
B('B) Reject AI Result: Sonographer disagrees. MUST manually annotate (draw) ≥1 region. Server validates annotation count before allowing Sign.')
P('')
P('Step 4 — Sign & Publish:', bold=True)
P('Sonographer fills image_description, professional_findings, conclusion → clicks "Ký & Công bố". System: saves ultrasound_report (versioned, is_current=1), updates test_order to Completed, SENDS NOTIFICATION to clinical doctor (critical handoff).')
P("BEGIN TRANSACTION;")
P("INSERT INTO ultrasound_reports (test_order_id, version, image_description, professional_findings, conclusion, report_status, is_current, created_by, signed_by_user_id, signed_name, signed_at, created_at)")
P("VALUES (?, ?, ?, ?, ?, 'FINAL', 1, ?, ?, ?, GETDATE(), GETDATE());")
P("UPDATE ultrasound_reports SET is_current = 0 WHERE test_order_id = ? AND id != ?;")
P("UPDATE test_orders SET status = 'Completed' WHERE id = ?;")
P("COMMIT;")
P("-- Then notify clinical doctor:")
P("INSERT INTO notifications (user_id, title, content, channel, is_read, created_at)")
P("SELECT u.id, N'Kết quả siêu âm đã sẵn sàng', ?, 'System', 0, GETDATE()")
P("FROM users u JOIN doctors d ON u.id = d.user_id JOIN test_orders tor ON d.id = tor.doctor_id WHERE tor.id = ?;")

doc.add_page_break()

# ── 6. DIAGNOSIS & PRESCRIPTION ──
H('6. Diagnosis & Prescription', 2)

H('6.1 UC-31_Diagnose and Prescribe', 3)
P('Doctor enters final diagnosis and prescribes medications. Ultrasound results displayed read-only above diagnosis form. Doctor CANNOT edit sonographer\'s conclusion.')
P('')
P('Business Rules:', bold=True)
B('Blocked if WAITING_ULTRASOUND stage: "Đang chờ kết quả siêu âm. Vui lòng đợi bác sĩ siêu âm hoàn tất."')
B('Blocked if WAITING_PAYMENT stage: "Bệnh nhân chưa thanh toán dịch vụ siêu âm."')
B('Doctor must confirm viewing ultrasound results before diagnosis entry unlocks')
B('Doctor CANNOT modify sonographer conclusion (UI: readonly textarea, Server: no update query for ultrasound_reports)')
B('Prescription purchase decision: AT_CLINIC → creates PRESCRIPTION invoice (Unpaid); EXTERNAL → no invoice')
B('Prescription does NOT block finalization (UC-32)')
P('')
P('SQL Commands:', bold=True)
P('1/ Save diagnosis:')
P("UPDATE medical_records SET final_diagnosis = ?, treatment_plan = ?, next_appointment_date = ?, referred_to = ?, updated_at = GETDATE(), updated_by = ? WHERE id = ?;")
P('2/ Create prescription with purchase decision:')
P("INSERT INTO prescriptions (medical_record_id, prescription_code, status, purchase_decision, purchase_decided_at, purchase_decided_by, created_at)")
P("VALUES (?, ?, 'issued', ?, GETDATE(), ?, GETDATE());")
P('3/ Insert prescription items (loop per medicine):')
P("INSERT INTO prescription_items (prescription_id, medicine_id, quantity, dosage) VALUES (?, ?, ?, ?);")
P('4/ If AT_CLINIC — create PRESCRIPTION invoice:')
P("INSERT INTO invoices (appointment_id, total_amount, status, invoice_type, created_at)")
P("SELECT a.id, ?, 'Unpaid', 'PRESCRIPTION', GETDATE() FROM medical_records mr JOIN appointments a ON mr.appointment_id = a.id WHERE mr.id = ?;")

H('6.2 UC-32_Finalize Medical Record', 3)
P('Doctor finalizes the medical record. SINGLE atomic transaction: medical_record status → Final, appointment status → SUCCESS, queue_number cleared, audit log. Prescription payment NOT required.')
P('')
P('Business Rules:', bold=True)
B('Requires: READY_TO_FINALIZE stage (diagnosis entered + prescription decision made + no pending ultrasound)')
B('Does NOT require: prescription invoice paid (handled separately by staff)')
B('Transaction rollback on any failure — medical_record and appointment stay in sync')
B('After finalization: entire medical record page becomes read-only for everyone')
B('Patient can now view the finalized record')
P('')
P('SQL Commands:', bold=True)
P('BEGIN TRANSACTION;')
P("UPDATE medical_records SET status = 'Final', updated_at = GETDATE() WHERE id = ? AND status = 'Draft';")
P("UPDATE appointments SET status = 'SUCCESS', queue_number = NULL WHERE id = ? AND status = 'InProgress';")
P("INSERT INTO audit_logs (user_id, action, table_name, new_value, ip_address, created_at) VALUES (?, 'FINALIZE_RECORD', 'medical_records', ?, ?, GETDATE());")
P('COMMIT;')
P('-- Notify patient:')
P("INSERT INTO notifications (user_id, title, content, channel, is_read, created_at)")
P("SELECT p.user_id, N'Kết quả khám đã sẵn sàng', N'Bạn có thể xem hồ sơ bệnh án trong mục Hồ sơ bệnh án.', 'System', 0, GETDATE()")
P("FROM appointments a JOIN patients p ON a.patient_id = p.id WHERE a.id = ?;")

doc.add_page_break()

# ── 7. PAYMENT MANAGEMENT ──
H('7. Payment Management', 2)
P('CAMS payment system is CASH-ONLY. No electronic payment gateway per project rules. All payments processed by Staff at reception counter.')

P('')
P('Three Invoice Types:', bold=True)
B('PRE_EXAM (Exam Fee): Created on booking approval. Amount = appointment.base_fee (locked at booking time, NEVER recalculated).')
B('POST_EXAM (Imaging Service): Created on ultrasound order. Amount = service.price at time of order (locked).')
B('PRESCRIPTION (Medication): Created on prescription with AT_CLINIC decision. Amount = sum of medicine prices × quantities.')

P('')
P('Invoice State Machine:', bold=True)
B('Unpaid → Paid: Staff confirms cash received. Records: paid_at=GETDATE(), paid_by_user_id, payment_method=CASH. Anti-duplicate check in WHERE status=\'Unpaid\'.')
B('Paid → Refunded: Staff processes refund. Records: refunded_at, refunded_by_user_id, refund_reason. If POST_EXAM → auto-cancels test_order.')
B('Unpaid → Cancelled: Auto-cancelled when appointment is cancelled or ultrasound order cancelled.')

P('')
P('Revenue Calculation (real-time, no aggregate table):', bold=True)
P("SELECT invoice_type, SUM(total_amount) FROM invoices WHERE status = 'Paid' AND confirmed_at BETWEEN ? AND ? GROUP BY invoice_type;")
P("-- Subtract refunds:")
P("SELECT invoice_type, SUM(total_amount) FROM invoices WHERE status = 'Refunded' AND refunded_at BETWEEN ? AND ? GROUP BY invoice_type;")

doc.add_page_break()

# ── 8. REVENUE & STATISTICS ──
H('8. Revenue & Statistics', 2)

H('8.1 UC-56_Manage Revenue', 3)
P('Manager views revenue report for any date range. Revenue = SUM(Paid invoices) - SUM(Refunded invoices). Broken down by PRE_EXAM, POST_EXAM, PRESCRIPTION. No pre-aggregated table — always queries invoices directly.')
P('')
P('SQL Commands:', bold=True)
P('1/ Total revenue by type (Paid - Refunded):')
P("SELECT i.invoice_type, ISNULL(SUM(CASE WHEN i.status='Paid' THEN i.total_amount ELSE 0 END), 0) - ISNULL(SUM(CASE WHEN i.status='Refunded' THEN i.total_amount ELSE 0 END), 0) AS net_revenue")
P("FROM invoices i WHERE (i.status = 'Paid' AND i.confirmed_at BETWEEN ? AND ?) OR (i.status = 'Refunded' AND i.refunded_at BETWEEN ? AND ?)")
P("GROUP BY i.invoice_type;")
P('2/ View Service Details (drill-down):')
P("SELECT i.id, i.invoice_type, i.total_amount, i.status, i.confirmed_at, i.paid_at, p.full_name AS patient_name, u.full_name AS staff_name")
P("FROM invoices i LEFT JOIN appointments a ON i.appointment_id = a.id LEFT JOIN patients p ON a.patient_id = p.id LEFT JOIN users u ON i.confirmed_by = u.id")
P("WHERE i.confirmed_at BETWEEN ? AND ? ORDER BY i.confirmed_at DESC;")

H('8.2 UC-57_View Service Statistics', 3)
P('Manager views service utilization statistics (COUNT of orders, not money). Shows: service name, number of times ordered, number of times completed. Filterable by date range.')
P('')
P('SQL Commands:', bold=True)
P('1/ Service order counts:')
P("SELECT s.service_name, COUNT(tor.id) AS total_ordered, SUM(CASE WHEN tor.status='Completed' THEN 1 ELSE 0 END) AS total_completed")
P("FROM test_orders tor JOIN services s ON tor.service_id = s.id WHERE tor.created_at BETWEEN ? AND ? GROUP BY s.service_name ORDER BY total_ordered DESC;")

doc.add_page_break()

# ── 9. SERVICE & PRICE MANAGEMENT ──
H('9. Service & Price Management', 2)

H('9.1 UC-47_Manage Medical Services', 3)
P('Manager/Admin manages medical services catalog. CRUD operations. Price changes create price_history entries. Cannot delete services that have been used — deactivate only. New price applies only to future orders; existing invoices unaffected.')
P('')
P('SQL Commands:', bold=True)
P('1/ List all services:')
P("SELECT s.*, sc.category_name FROM services s LEFT JOIN service_categories sc ON s.category_id = sc.id ORDER BY s.is_active DESC, s.service_name;")
P('2/ Create service:')
P("INSERT INTO services (service_name, price, service_code, category_id, duration_mins, requires_fasting, requires_full_bladder, description, is_active, created_at)")
P("VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, GETDATE());")
P('3/ Update service with price change (transactional with price_history):')
P("BEGIN TRANSACTION;")
P("DECLARE @oldPrice DECIMAL(18,2); SELECT @oldPrice = price FROM services WHERE id = ?;")
P("UPDATE services SET service_name = ?, price = ?, description = ?, is_active = ?, updated_at = GETDATE() WHERE id = ?;")
P("IF @oldPrice != ? INSERT INTO price_history (service_id, old_price, new_price, change_reason, changed_by, created_at) VALUES (?, @oldPrice, ?, ?, ?, GETDATE());")
P("COMMIT;")
P('4/ View price history:')
P("SELECT ph.*, s.service_name, u.full_name AS changed_by_name FROM price_history ph JOIN services s ON ph.service_id = s.id LEFT JOIN users u ON ph.changed_by = u.id WHERE ph.service_id = ? ORDER BY ph.created_at DESC;")

H('9.2 UC-48_View Price Adjustment History', 3)
P('Manager views all price changes for a service or date range. Shows: service name, old price, new price, difference, changed by, timestamp.')
P('')
P('SQL Commands:', bold=True)
P("SELECT ph.*, s.service_name, u.full_name AS changed_by_name, (ph.new_price - ISNULL(ph.old_price, 0)) AS difference")
P("FROM price_history ph JOIN services s ON ph.service_id = s.id LEFT JOIN users u ON ph.changed_by = u.id")
P("WHERE ph.created_at BETWEEN ? AND ? ORDER BY ph.created_at DESC;")

doc.add_page_break()

# ── 10. USER, ROLE & SECURITY ──
H('10. User, Role & Security Management', 2)

H('10.1 UC-60_Manage Users', 3)
P('Admin manages all user accounts. List with pagination, search by username/email, filter by role/status/date range. Create staff accounts (any role). Update account status (ACTIVE/INACTIVE/LOCKED). Soft delete (is_deleted=1).')
P('')
P('SQL Commands:', bold=True)
P("SELECT u.*, r.role_name FROM users u LEFT JOIN roles r ON u.role_id = r.id WHERE u.is_deleted = 0 ORDER BY u.created_at DESC OFFSET ? ROWS FETCH NEXT ? ROWS ONLY;")
P("INSERT INTO users (username, full_name, email, password_hash, phone, role_id, status, is_verified, auth_provider, is_deleted, created_at) VALUES (?, ?, ?, ?, ?, ?, 'ACTIVE', 1, 'LOCAL', 0, GETDATE());")
P("UPDATE users SET status = ? WHERE id = ?;")

H('10.2 UC-61_Manage Roles & Permissions', 3)
P('Admin manages RBAC: roles and their permissions. Changes bump global permissions version → AuthorizationFilter reloads on next request.')
P('')
P('SQL Commands:', bold=True)
P("SELECT * FROM roles; SELECT * FROM permissions; SELECT rp.* FROM role_permissions rp;")
P("INSERT INTO role_permissions (role_id, permission_id) VALUES (?, ?); DELETE FROM role_permissions WHERE role_id = ? AND permission_id = ?;")

H('10.3 UC-63_View Audit Logs', 3)
P('Admin views system audit trail. Filterable by user, action, table, date range.')
P('')
P('SQL Commands:', bold=True)
P("SELECT al.*, u.full_name AS user_name FROM audit_logs al LEFT JOIN users u ON al.user_id = u.id WHERE al.created_at BETWEEN ? AND ? ORDER BY al.created_at DESC;")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# BUSINESS RULES SUMMARY
# ═══════════════════════════════════════════════════════════════
H('Appendix: Business Rules Summary', 1)
P('All finalized business decisions for CAMS:')
T(['#', 'Issue', 'Decision'],
  [
      ['1', 'Payment method', 'CASH ONLY at reception counter. NO electronic payment gateway.'],
      ['2', 'Patient self-marking paid', 'ABSOLUTELY NOT. Only Staff confirms payment.'],
      ['3', 'Change appointment', 'Only change slot WITHIN SAME DOCTOR. Max 2 times (counted from audit log). Change doctor → cancel + rebook.'],
      ['4', 'Cancel/reschedule deadline', '2 hours before shift start time.'],
      ['5', 'Cancel when paid', 'Patient cannot self-cancel. Redirect to reception for refund processing.'],
      ['6', 'Prescription blocks finalization?', 'NO. Doctor finalizes; medication invoice handled separately by Staff.'],
      ['7', 'Multiple ultrasound orders', 'ALLOWED. Diagnosis unlocks only when ALL orders completed.'],
      ['8', 'AI Engine failure', 'ALLOW skip AI. Manual draw entirely. Log "ca không qua AI".'],
      ['9', 'Statistics vs Revenue', 'SEPARATE screens. Statistics = COUNT of services. Revenue = MONEY (Paid - Refunded).'],
      ['10', 'Service price change', 'Existing invoices UNCHANGED. New price only for future orders. Recorded in price_history.'],
      ['11', 'Overnight appointments', 'ACCEPTED. No auto-cancel. Manager has "stuck >24h" monitoring list.'],
      ['12', 'Parallel consultations', 'ALLOWED. Doctor can have multiple InProgress appointments simultaneously.'],
      ['13', 'Exam fee locking', 'base_fee calculated ONCE at booking (200,000 + experience_years×50,000). Stored in appointment. NEVER recalculated.'],
      ['14', 'Invoice amount source', 'PRE_EXAM copies base_fee. POST_EXAM copies service.price at order time. PRESCRIPTION sums medicine prices.'],
      ['15', 'Queue ordering', 'By priority flag → status score → created_at. Renumbered on check-in/priority change.'],
      ['16', 'Auto-move patients', 'When shift ends, waiting patients auto-moved to doctor\'s next available shift. Transactional + notify.'],
      ['17', 'Doctor edit sonographer conclusion', 'NOT ALLOWED. Read-only display of ultrasound report.'],
      ['18', 'Staff edit diagnosis/prescription', 'NOT ALLOWED. Only doctor can enter diagnosis and prescribe.'],
      ['19', 'No permanent stuck state', 'Cancel Ultrasound Order UC provides escape hatch. Manager monitors >24h stuck cases.'],
      ['20', 'Anti-duplicate (idempotency)', 'Double-click/F5 on any action: second request rejected. booked_count never double-modified.'],
      ['21', 'Transaction for multi-table writes', 'All status changes involving money or multiple tables wrapped in BEGIN/COMMIT/ROLLBACK.'],
      ['22', 'Audit log', 'Every status change and financial operation logged: user, action, table, old_value, new_value, IP, timestamp.'],
      ['23', 'Ultrasound image storage', 'Raw bytes stored as-is. NO ImageIO re-encode (preserves diagnostic quality). Only avatars re-encoded.'],
      ['24', 'Server-side validation', 'ALL business validations checked at SERVER. JSP/JS hiding/disabling is cosmetic layer only.'],
      ['25', 'Ownership check', 'Every action receiving ID from URL/form verifies ownership/assignment. ID swapping → blocked + Vietnamese error.'],
  ],
  [0.5, 5, 11.5])

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
output_path = r'c:\Users\admin\IdeaProjects\Clinic-Appointment-Management-System\docs\CAMS_RDS_V2.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print('Done!')
